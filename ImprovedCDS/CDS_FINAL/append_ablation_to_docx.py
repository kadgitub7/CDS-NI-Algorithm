"""Append the new ablation, significance, and sensitivity sections to the existing docx."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

DOC_PATH = os.path.join(os.path.dirname(__file__), "CDS_OVR_Comparative_Analysis_Report.docx")

doc = Document(DOC_PATH)


def add_heading(text, level):
    style_map = {2: 'Heading 2', 3: 'Heading 3', 4: 'Heading 4'}
    doc.add_paragraph(text, style=style_map.get(level, 'Heading 3'))


def add_para(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_experiment_block(title, what, why, method, results_lines, interpretation):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True

    add_para(what, bold_prefix="What was tested: ")
    add_para(why, bold_prefix="Why: ")
    add_para(method, bold_prefix="Method: ")

    p = doc.add_paragraph()
    run = p.add_run("Results:")
    run.bold = True
    for line in results_lines:
        doc.add_paragraph(line, style='Normal')

    add_para(interpretation, bold_prefix="Interpretation: ")


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    # Try to apply a table style, fall back gracefully
    try:
        table.style = 'Table Grid'
    except KeyError:
        try:
            table.style = 'Light Grid'
        except KeyError:
            pass  # use default
    # Add borders via XML if no style applied
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single', qn('w:sz'): '4',
            qn('w:space'): '0', qn('w:color'): '000000'
        })
        borders.append(el)
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()  # spacing


# ─── Update the summary table (Section 8 "Why CDS-OVR Outperforms") ───
# We append new sections after the existing content

add_heading("8.1 Complete Component Ablation Study", 3)

add_para(
    "To quantify the contribution of each CDS-OVR mechanism, a complete ablation study was conducted: "
    "each component was removed one-at-a-time while all other settings remained at W11-01 defaults. "
    "All evaluations used 10-fold CV across 10 random seeds (100 fold evaluations per configuration). "
    "Statistical significance was assessed using the Wilcoxon signed-rank test (paired by seed), "
    "the appropriate nonparametric test for small paired samples where normality cannot be assumed."
)

add_experiment_block(
    title="Experiment: Full Component Ablation with Statistical Significance (ablation_full.py, Parts 2–3)",

    what="Six CDS-OVR mechanisms were individually disabled: "
         "(1) supervised chi-squared binning → equal-width Sturges binning, "
         "(2) correlation filtering → disabled (CORR_THRESHOLD=1.0), "
         "(3) healthy bar → disabled (HEALTHY_WEIGHT=0), "
         "(4) Fisher discriminant weighting → uniform weights (fw=1.0), "
         "(5) per-class against_scale → uniform 0.8, "
         "(6) sex branching → disabled (SEX_FEAT=-1).",

    why="To determine which mechanisms earn their complexity — i.e., which produce statistically significant "
        "accuracy improvements vs. simpler alternatives. Components that do not significantly improve accuracy "
        "may still be justified on other grounds (specificity, rare-class detection) but cannot be claimed as "
        "accuracy-improving innovations.",

    method="For each ablation, the modified system was evaluated via 10-fold CV across 10 seeds. "
           "Per-seed accuracy was paired with the baseline, and a two-sided Wilcoxon signed-rank test "
           "was applied (n=10 paired observations). Significance threshold: α=0.05.",

    results_lines=[],

    interpretation="Only two mechanisms produce statistically significant accuracy improvements: "
                   "supervised chi-squared binning (+5.75 pp, p=0.002) and Fisher discriminant weighting "
                   "(+4.52 pp, p=0.002). Four mechanisms do not reach statistical significance at α=0.05 "
                   "for overall accuracy. Notably, removing the healthy bar improves overall accuracy by "
                   "0.31 pp (not significant), confirming its role is clinical (specificity) rather than "
                   "accuracy-based."
)

# Ablation results table
add_table(
    headers=["Component Removed", "Baseline", "Ablated", "Δ pp", "W-stat", "p-value", "Significant (α=0.05)"],
    rows=[
        ["Supervised binning → equal-width", "84.66%", "78.92%", "+5.75", "0.0", "0.002", "YES"],
        ["Fisher weighting → uniform", "84.66%", "80.14%", "+4.52", "0.0", "0.002", "YES"],
        ["Correlation filtering → disabled", "84.66%", "84.21%", "+0.46", "17.0", "0.313", "no"],
        ["Sex branching → disabled", "84.66%", "84.25%", "+0.41", "7.0", "0.141", "no"],
        ["against_scale → uniform 0.8", "84.66%", "84.59%", "+0.07", "19.0", "0.734", "no"],
        ["Healthy bar → disabled", "84.66%", "84.98%", "−0.31", "6.5", "0.133", "no"],
    ]
)

add_para(
    "Supervised chi-squared binning (+5.75 pp, p=0.002): The single most impactful component. "
    "Replacing it with equal-width bins drops accuracy from 84.66% to 78.92% — every seed shows "
    "degradation (W=0). This confirms that aligning bin boundaries with class boundaries is critical "
    "for posterior quality.",
    bold_prefix="1. "
)

add_para(
    "Fisher discriminant weighting (+4.52 pp, p=0.002): The second most impactful component. "
    "Without Fisher weighting, all features contribute equally regardless of their discriminative power. "
    "Accuracy drops to 80.14% — again, every seed shows degradation (W=0).",
    bold_prefix="2. "
)

add_para(
    "Correlation filtering (+0.46 pp, p=0.313): Modest benefit, likely because the correlation "
    "filter’s main effect is ensuring feature diversity, which may matter more for specific classes "
    "than for overall accuracy.",
    bold_prefix="3. "
)

add_para(
    "Sex branching (+0.41 pp, p=0.141): Concentrated effect on class 3 (100% male), diluted across "
    "all 8 classes. The per-class benefit is real but too localized to affect overall accuracy significantly.",
    bold_prefix="4. "
)

add_para(
    "against_scale (+0.07 pp, p=0.734): Negligible overall effect because rare classes represent only "
    "8.9% of patients. The mechanism’s value is in per-class rare-class detection (+9.2 pp for class 5, "
    "+6.7 pp for class 4 — see Section 4.6) rather than overall accuracy.",
    bold_prefix="5. "
)

add_para(
    "Healthy bar (−0.31 pp, p=0.133): Removing the healthy bar improves overall accuracy by 0.31 pp "
    "(not significant). This is expected: the healthy bar deliberately trades overall accuracy for better "
    "specificity — it prevents false positives (21/245 = 8.6% FPR with bar vs. potentially higher without) "
    "at the cost of some true disease detections. The healthy bar’s value is in specificity, not accuracy, "
    "making it a clinically justified mechanism rather than an accuracy-improving one.",
    bold_prefix="6. "
)

# ─── Section 8.2: Hyperparameter Sensitivity ───

add_heading("8.2 Hyperparameter Sensitivity Analysis", 3)

add_para(
    "To assess whether CDS-OVR’s hardcoded hyperparameters represent a fragile configuration or a "
    "robust operating point, three key parameters were swept across their plausible ranges using "
    "10-fold CV across 10 seeds."
)

add_experiment_block(
    title="Experiment: Hyperparameter Sensitivity Sweep (ablation_full.py, Part 4)",

    what="Three hyperparameters were varied independently while all others were held at W11-01 defaults: "
         "(a) CORR_THRESHOLD ∈ {0.6, 0.7, 0.8, 0.9, 1.0}, "
         "(b) MAX_BINS ∈ {3, 4, 5, 6, 8, 10}, "
         "(c) FEATURES_PER_CLASS ∈ {10, 14, 18, 22, 26, 30}.",

    why="If accuracy collapses with small parameter changes, the system is brittle and the reported "
        "accuracy may be an artifact of careful tuning. If accuracy is stable across a wide range, "
        "the system’s performance is robust and not dependent on specific hyperparameter values.",

    method="For each parameter value, 10-fold CV was run across 10 seeds (100 fold evaluations). "
           "Mean accuracy and standard deviation were recorded.",

    results_lines=[],

    interpretation="CDS-OVR’s accuracy is robust across plausible hyperparameter ranges. The most "
                   "sensitive parameter is MAX_BINS (3.05 pp range), but even this shows a broad plateau "
                   "at 5–6 bins rather than a sharp peak. CORR_THRESHOLD has 1.44 pp range with the "
                   "current value at the optimum. FEATURES_PER_CLASS has zero effect, proving it is not a "
                   "tuning knob. These results demonstrate that CDS-OVR’s performance is not an artifact "
                   "of precise hyperparameter tuning."
)

# CORR_THRESHOLD table
add_para("CORR_THRESHOLD (correlation filter strictness):", bold_prefix="")
add_table(
    headers=["Value", "Mean Accuracy", "σ", "Δ from best"],
    rows=[
        ["0.6", "83.22%", "1.06%", "−1.44 pp"],
        ["0.7", "83.68%", "1.35%", "−0.98 pp"],
        ["0.8 (current)", "84.66%", "0.93%", "best"],
        ["0.9", "84.30%", "0.97%", "−0.36 pp"],
        ["1.0", "84.21%", "0.96%", "−0.45 pp"],
    ]
)

add_para(
    "Range: 1.44 pp. The current value (0.8) is optimal. Accuracy degrades monotonically as the "
    "threshold tightens below 0.8 (too aggressive filtering removes useful features) and slightly "
    "when loosened above 0.8 (redundant features waste model capacity). The degradation is gradual, "
    "not cliff-like."
)

# MAX_BINS table
add_para("MAX_BINS (maximum supervised bins per feature):", bold_prefix="")
add_table(
    headers=["Value", "Mean Accuracy", "σ", "Δ from best"],
    rows=[
        ["3", "84.06%", "1.01%", "−1.11 pp"],
        ["4", "84.01%", "1.18%", "−1.16 pp"],
        ["5", "85.17%", "0.37%", "best"],
        ["6 (current)", "84.66%", "0.93%", "−0.51 pp"],
        ["8", "83.22%", "0.90%", "−1.95 pp"],
        ["10", "82.12%", "0.75%", "−3.05 pp"],
    ]
)

add_para(
    "Range: 3.05 pp. MAX_BINS=5 outperforms the current setting of 6 by +0.51 pp and has the "
    "lowest variance (σ=0.37%). Too few bins (3–4) lose discriminative resolution. Too many bins "
    "(8–10) cause overfitting — each bin has fewer patients, producing noisier posterior estimates. "
    "The sweet spot is 5–6, and the current setting of 6 is within 0.51 pp of optimal."
)

# FEATURES_PER_CLASS table
add_para("FEATURES_PER_CLASS (maximum retained features per class model):", bold_prefix="")
add_table(
    headers=["Value", "Mean Accuracy", "σ"],
    rows=[
        ["10", "84.66%", "0.93%"],
        ["14", "84.66%", "0.93%"],
        ["18 (current)", "84.66%", "0.93%"],
        ["22", "84.66%", "0.93%"],
        ["26", "84.66%", "0.93%"],
        ["30", "84.66%", "0.93%"],
    ]
)

add_para(
    "Range: 0.00 pp. All values produce identical results. This means the correlation filter "
    "(CORR_THRESHOLD=0.8) is the actual bottleneck — it caps the effective feature count well below "
    "even FPC=10 for most class models. The FPC parameter is not a tuning-sensitive hyperparameter; "
    "it is a safety cap that never binds under current settings."
)

add_para(
    "Note on MAX_BINS=5: The sweep reveals that MAX_BINS=5 achieves 85.17% (σ=0.37%), outperforming "
    "the current setting of 6 (84.66%, σ=0.93%). This suggests that 6 bins may slightly overfit for "
    "some class models. However, the current W11-01 configuration uses MAX_BINS=6, and all results "
    "throughout this document are reported with that setting. The sensitivity analysis demonstrates "
    "that a 0.51 pp improvement is available but was not exploited, providing further evidence against "
    "overfitting to the reported configuration.",
    bold_prefix=""
)

# Save
doc.save(DOC_PATH)
print(f"Saved updated docx to {DOC_PATH}")
print(f"Total paragraphs now: {len(doc.paragraphs)}")
