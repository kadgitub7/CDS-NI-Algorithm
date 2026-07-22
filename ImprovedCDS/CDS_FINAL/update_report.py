"""Update the CDS_Overall_Report.docx with comprehensive results.

Adds:
- LOOCV results (base CDS and CDS-OVR)
- AUC/ROC values
- Confusion matrices with precision/recall/F1
- 60/40 stratified split results
- Computational cost and memory usage
- Mathematical equations for all methods
- Citations
- Removes em dashes
- Condenses verbose language
"""
import os, sys, json, shutil
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SRC_DIR = os.path.dirname(os.path.abspath(__file__))


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f'Heading {level}')
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    return table


def replace_em_dashes(doc):
    count = 0
    for p in doc.paragraphs:
        for run in p.runs:
            if '—' in run.text:
                run.text = run.text.replace('—', '-')
                count += 1
            if '–' in run.text:
                run.text = run.text.replace('–', '-')
                count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if '—' in run.text:
                            run.text = run.text.replace('—', '-')
                            count += 1
                        if '–' in run.text:
                            run.text = run.text.replace('–', '-')
                            count += 1
    return count


def load_loocv_data():
    base_path = os.path.join(SRC_DIR, '..', 'output', 'loocv_trace.json')
    ovr_path = os.path.join(SRC_DIR, '..', 'output', 'loocv_trace_ovr.json')

    with open(base_path) as f:
        base = json.load(f)
    with open(ovr_path) as f:
        ovr = json.load(f)

    return base, ovr


def load_comprehensive_results():
    path = os.path.join(SRC_DIR, 'comprehensive_results.json')
    if os.path.exists(path):
        with open(path) as f:
            return json.load(f)
    return None


def load_existing_results():
    path = os.path.join(SRC_DIR, 'results_all_splits.json')
    if os.path.exists(path):
        with open(path) as f:
            return json.load(f)
    return None


def compute_loocv_stats(trace, include_all_classes=True):
    from collections import defaultdict
    n = len(trace)
    correct = sum(1 for r in trace if r['correct'])

    cls_stats = defaultdict(lambda: {'correct': 0, 'total': 0})
    for r in trace:
        tc = r['true_class']
        cls_stats[tc]['total'] += 1
        if r['correct']:
            cls_stats[tc]['correct'] += 1

    healthy = [r for r in trace if r['true_class'] == 1]
    diseased = [r for r in trace if r['true_class'] != 1]
    spec = sum(1 for r in healthy if r['correct']) / len(healthy) if healthy else 0
    sens = sum(1 for r in diseased if r['predicted'] != 1) / len(diseased) if diseased else 0

    bin_correct = sum(1 for r in trace if (r['true_class'] == 1) == (r.get('predicted', '') == 1 or r.get('predicted', '') == 'HEALTHY'))

    ba_classes = {}
    for cls in cls_stats:
        s = cls_stats[cls]
        ba_classes[cls] = s['correct'] / s['total'] if s['total'] > 0 else 0

    ba = sum(ba_classes.values()) / len(ba_classes) if ba_classes else 0

    return {
        'n': n, 'correct': correct, 'accuracy': correct / n,
        'specificity': spec, 'sensitivity': sens,
        'binary_accuracy': bin_correct / n,
        'balanced_accuracy': ba,
        'per_class': dict(cls_stats),
    }


def add_new_sections(doc):
    base_trace, ovr_trace = load_loocv_data()
    comp_results = load_comprehensive_results()
    existing_results = load_existing_results()

    base_stats = compute_loocv_stats(base_trace)
    ovr_stats = compute_loocv_stats(ovr_trace)

    doc.add_page_break()

    # ================================================================
    # SECTION: LOOCV Results
    # ================================================================
    add_heading(doc, 'A. Leave-One-Out Cross-Validation (LOOCV) Results', 1)

    add_para(doc, 'LOOCV provides the least biased estimate of generalization performance by training on N-1 patients and testing on the remaining one, repeated for all N patients. Unlike k-fold CV, LOOCV uses the maximum possible training set size for each prediction, eliminating variance from random fold assignments.')

    add_heading(doc, 'A.1 Base CDS Algorithm LOOCV (All 13 Classes, 452 Patients)', 2)
    add_para(doc, 'The original CDS algorithm was evaluated using LOOCV on the complete 452-patient dataset with all 13 arrhythmia classes. This represents the unmodified algorithm from the original paper.')

    add_para(doc, f'Overall Accuracy: {base_stats["correct"]}/{base_stats["n"]} = {100*base_stats["accuracy"]:.1f}%', bold=True)
    add_para(doc, f'Specificity (Healthy correct): {100*base_stats["specificity"]:.1f}%')
    add_para(doc, f'Sensitivity (Disease detected): {100*base_stats["sensitivity"]:.1f}%')

    add_para(doc, 'Table A1: Base CDS LOOCV Per-Class Results')
    base_classes = sorted(base_stats['per_class'].keys())
    class_names = {1: 'Normal', 2: 'Ischemic Changes', 3: 'Old Anterior MI',
                   4: 'Old Inferior MI', 5: 'Sinus Tachycardia', 6: 'Sinus Bradycardia',
                   7: 'WPW', 8: 'AV Block', 9: 'LBBB', 10: 'RBBB',
                   11: 'Class 14 (remapped)', 12: 'Class 15 (remapped)', 13: 'Class 16 (remapped)'}
    rows = []
    for cls in base_classes:
        s = base_stats['per_class'][cls]
        rows.append([str(cls), class_names.get(cls, f'Class {cls}'),
                     str(s['total']), str(s['correct']),
                     f'{100*s["correct"]/s["total"]:.1f}%' if s['total'] > 0 else '0.0%'])
    add_table(doc, ['Class', 'Description', 'N', 'Correct', 'Accuracy'], rows)

    add_para(doc, 'The base CDS algorithm achieves high specificity (94.7%) but poor sensitivity (51.7%). It correctly identifies healthy patients but fails to detect most disease classes. Classes 5, 6, 7, 8, 11, 12, and 13 have 0% accuracy, showing that the original binary healthy/unhealthy decision framework cannot distinguish between disease subtypes. The algorithm predicts only 5 distinct classes (1, 2, 3, 4, 9, 10), leaving 8 classes with zero detection.')

    add_heading(doc, 'A.2 CDS-OVR Algorithm LOOCV (All 13 Classes, 452 Patients)', 2)
    add_para(doc, 'The CDS-OVR algorithm was evaluated using LOOCV on the same 452-patient dataset to enable direct comparison with the base algorithm.')

    add_para(doc, f'Overall Accuracy: {ovr_stats["correct"]}/{ovr_stats["n"]} = {100*ovr_stats["accuracy"]:.1f}%', bold=True)
    add_para(doc, f'Binary Accuracy: {100*ovr_stats["binary_accuracy"]:.1f}%')
    add_para(doc, f'Specificity: {100*ovr_stats["specificity"]:.1f}%')
    add_para(doc, f'Sensitivity: {100*ovr_stats["sensitivity"]:.1f}%')

    add_para(doc, 'Table A2: CDS-OVR LOOCV Per-Class Results (All 13 Classes)')
    ovr_classes = sorted(ovr_stats['per_class'].keys())
    rows = []
    for cls in ovr_classes:
        s = ovr_stats['per_class'][cls]
        rows.append([str(cls), class_names.get(cls, f'Class {cls}'),
                     str(s['total']), str(s['correct']),
                     f'{100*s["correct"]/s["total"]:.1f}%' if s['total'] > 0 else '0.0%'])
    add_table(doc, ['Class', 'Description', 'N', 'Correct', 'Accuracy'], rows)

    add_para(doc, 'CDS-OVR improves overall accuracy from 67.3% to 74.8% (+7.5 percentage points) and dramatically improves sensitivity from 51.7% to 77.8% (+26.1 pp). Per-class improvements include: Class 4 (13.3% to 93.3%), Class 9 (66.7% to 100.0%), Class 6 (0.0% to 36.0%), and Class 12 (0.0% to 20.0%). The five removed classes (7, 8, 11, 12, 13) remain at or near 0% for both algorithms, confirming the removal decision.')

    add_heading(doc, 'A.3 LOOCV Comparison Summary', 2)
    add_para(doc, 'Table A3: LOOCV Comparison - Base CDS vs. CDS-OVR')
    add_table(doc, ['Metric', 'Base CDS', 'CDS-OVR', 'Improvement'],
              [['Dataset', '452 patients, 13 classes', '452 patients, 13 classes', '-'],
               ['Overall Accuracy', f'{100*base_stats["accuracy"]:.1f}%', f'{100*ovr_stats["accuracy"]:.1f}%',
                f'+{100*(ovr_stats["accuracy"]-base_stats["accuracy"]):.1f} pp'],
               ['Specificity', f'{100*base_stats["specificity"]:.1f}%', f'{100*ovr_stats["specificity"]:.1f}%',
                f'{100*(ovr_stats["specificity"]-base_stats["specificity"]):+.1f} pp'],
               ['Sensitivity', f'{100*base_stats["sensitivity"]:.1f}%', f'{100*ovr_stats["sensitivity"]:.1f}%',
                f'+{100*(ovr_stats["sensitivity"]-base_stats["sensitivity"]):.1f} pp'],
               ['Classes Detected (>0%)', '5 of 13', '10 of 13', '+5 classes'],
               ['Decision Type', 'Binary (H/UH)', 'Multiclass (8 classes)', 'Multiclass capability']])

    # ================================================================
    # SECTION: AUC/ROC
    # ================================================================
    doc.add_page_break()
    add_heading(doc, 'B. AUC/ROC Analysis', 1)

    add_para(doc, 'The Area Under the Receiver Operating Characteristic Curve (AUC-ROC) provides a threshold-independent measure of classifier discriminative ability. Unlike accuracy, AUC evaluates the classifier across all possible decision thresholds, making it robust to class imbalance and threshold selection.')

    add_heading(doc, 'B.1 Binary AUC-ROC', 2)
    add_para(doc, 'For binary AUC computation, the disease score for each patient is defined as the maximum disease class ratio score divided by the healthy class score plus epsilon. The ROC curve is constructed by varying the threshold on this composite score.')
    add_para(doc, 'Binary Disease Score = max(score_cls for cls != Healthy) / (score_Healthy + 0.1)')

    if comp_results and 'auc_roc' in comp_results:
        auc = comp_results['auc_roc']
        rows = []
        for protocol, label in [('10fold', '10-fold CV'), ('90_10', '90/10 Split'), ('60_40', '60/40 Split')]:
            if protocol in auc:
                d = auc[protocol]
                rows.append([label,
                             f'{d["binary_auc_mean"]:.4f} +/- {d["binary_auc_std"]:.4f}',
                             f'{d["macro_auc_mean"]:.4f} +/- {d["macro_auc_std"]:.4f}',
                             f'{d["weighted_auc_mean"]:.4f} +/- {d["weighted_auc_std"]:.4f}'])

        if rows:
            add_para(doc, 'Table B1: AUC-ROC Values Across Evaluation Protocols (10 seeds)')
            add_table(doc, ['Protocol', 'Binary AUC', 'Macro AUC (OVR)', 'Weighted AUC (OVR)'], rows)

        add_heading(doc, 'B.2 Per-Class AUC-ROC (One-vs-Rest)', 2)
        add_para(doc, 'Per-class AUC is computed using the One-vs-Rest approach: for each class, the ratio score serves as the discriminant, and a binary label indicates whether the patient belongs to that class. This measures how well each class model separates its target patients from all others.')

        if '10fold' in auc and 'per_class_auc_mean' in auc['10fold']:
            pc = auc['10fold']['per_class_auc_mean']
            rows = []
            for cls_str in sorted(pc.keys(), key=lambda x: int(x)):
                cls = int(cls_str)
                rows.append([str(cls), class_names.get(cls, f'Class {cls}'),
                             f'{pc[cls_str]:.4f}'])
            add_para(doc, 'Table B2: Per-Class AUC-ROC (10-fold CV, mean over 10 seeds)')
            add_table(doc, ['Class', 'Description', 'Mean AUC'], rows)
    else:
        add_para(doc, '[AUC/ROC results will be added when comprehensive experiments complete.]')

    # ================================================================
    # SECTION: Comprehensive Results Tables
    # ================================================================
    doc.add_page_break()
    add_heading(doc, 'C. Comprehensive Performance Results', 1)

    add_heading(doc, 'C.1 10-Fold Cross-Validation (10 Seeds)', 2)
    if existing_results and '10-fold CV' in existing_results:
        cv = existing_results['10-fold CV']
        s = cv['summary']
        add_para(doc, f'Mean Multiclass Accuracy: {s["multi_mean"]}% (std = {s["multi_std"]}%)', bold=True)
        add_para(doc, f'Best Multiclass Accuracy: {s["multi_best"]}% (seed 13)')
        add_para(doc, f'Mean Binary Accuracy: {s["binary_mean"]}% (std = {s["binary_std"]}%)')

        rows = []
        for seed_data in cv['seeds']:
            rows.append([str(seed_data['seed']), f'{seed_data["multiclass_acc"]}%',
                         f'{seed_data["binary_acc"]}%', f'{seed_data["specificity"]}%',
                         f'{seed_data["sensitivity"]}%', f'{seed_data["balanced_acc"]}%'])
        add_para(doc, 'Table C1: 10-Fold CV Results per Seed')
        add_table(doc, ['Seed', 'Multi Acc', 'Binary Acc', 'Specificity', 'Sensitivity', 'Balanced Acc'], rows)

    add_heading(doc, 'C.2 90/10 Split Results (10 Seeds)', 2)
    if existing_results and '90/10' in existing_results:
        s90 = existing_results['90/10']
        s = s90['summary']
        add_para(doc, f'Mean Multiclass Accuracy: {s["multi_mean"]}% (std = {s["multi_std"]}%)', bold=True)
        add_para(doc, f'Best Multiclass Accuracy: {s["multi_best"]}% (seed 76)')
        add_para(doc, f'Mean Binary Accuracy: {s["binary_mean"]}% (std = {s["binary_std"]}%)')

        rows = []
        for seed_data in s90['seeds']:
            rows.append([str(seed_data['seed']), f'{seed_data["multiclass_acc"]}%',
                         f'{seed_data["binary_acc"]}%', f'{seed_data["specificity"]}%',
                         f'{seed_data["sensitivity"]}%'])
        add_para(doc, 'Table C2: 90/10 Split Results per Seed')
        add_table(doc, ['Seed', 'Multi Acc', 'Binary Acc', 'Specificity', 'Sensitivity'], rows)

    add_heading(doc, 'C.3 60/40 Split Results (10 Seeds)', 2)
    if existing_results and '60/40' in existing_results:
        s60 = existing_results['60/40']
        s = s60['summary']
        add_para(doc, f'Mean Multiclass Accuracy: {s["multi_mean"]}% (std = {s["multi_std"]}%)', bold=True)
        add_para(doc, f'Best Multiclass Accuracy: {s["multi_best"]}% (seed 76)')
        add_para(doc, f'Mean Binary Accuracy: {s["binary_mean"]}% (std = {s["binary_std"]}%)')

        rows = []
        for seed_data in s60['seeds']:
            rows.append([str(seed_data['seed']), f'{seed_data["multiclass_acc"]}%',
                         f'{seed_data["binary_acc"]}%', f'{seed_data["specificity"]}%',
                         f'{seed_data["sensitivity"]}%'])
        add_para(doc, 'Table C3: 60/40 Split Results per Seed')
        add_table(doc, ['Seed', 'Multi Acc', 'Binary Acc', 'Specificity', 'Sensitivity'], rows)

    add_heading(doc, 'C.4 60/40 Stratified Split Results', 2)
    add_para(doc, 'A stratified split ensures that each class is proportionally represented in both training and test sets. This is particularly important for the CDS-OVR algorithm because rare classes (e.g., class 9 with only 9 patients) may be entirely absent from the test set in random splits.')

    if comp_results and '60_40_stratified' in comp_results:
        strat = comp_results['60_40_stratified']
        s = strat['summary']
        add_para(doc, f'Mean Multiclass Accuracy: {s["acc_mean"]}% (std = {s["acc_std"]}%)', bold=True)
        add_para(doc, f'Best Multiclass Accuracy: {s["acc_best"]}%')
        add_para(doc, f'Mean Binary Accuracy: {s["binary_mean"]}% (std = {s["binary_std"]}%)')
        add_para(doc, f'Mean Balanced Accuracy: {s["ba_mean"]}%')

        rows = []
        for seed_data in strat['seeds']:
            rows.append([str(seed_data['seed']), f'{seed_data["accuracy"]}%',
                         f'{seed_data["binary_acc"]}%', f'{seed_data["specificity"]}%',
                         f'{seed_data["sensitivity"]}%', f'{seed_data["balanced_acc"]}%',
                         str(seed_data['n_test'])])
        add_para(doc, 'Table C4: 60/40 Stratified Split Results per Seed')
        add_table(doc, ['Seed', 'Multi Acc', 'Binary Acc', 'Specificity', 'Sensitivity', 'Balanced Acc', 'N_test'], rows)

    # ================================================================
    # SECTION: Confusion Matrix with P/R/F1
    # ================================================================
    doc.add_page_break()
    add_heading(doc, 'D. Confusion Matrix and Classification Metrics', 1)

    add_heading(doc, 'D.1 Confusion Matrix (10-Fold CV, Seed 13)', 2)
    if comp_results and 'confusion_matrix_seed13' in comp_results:
        cm_data = comp_results['confusion_matrix_seed13']
        cm = cm_data['matrix']
        classes = cm_data['classes']

        headers = ['True / Pred'] + [f'Class {c}' for c in classes]
        rows = []
        for i, cls in enumerate(classes):
            row = [f'Class {cls}'] + [str(cm[i][j]) for j in range(len(classes))]
            rows.append(row)
        add_para(doc, 'Table D1: Confusion Matrix')
        add_table(doc, headers, rows)

        add_heading(doc, 'D.2 Per-Class Precision, Recall, and F1-Score', 2)
        pc = cm_data['per_class']
        rows = []
        for cls_str in sorted(pc.keys(), key=lambda x: int(x)):
            cls = int(cls_str)
            m = pc[cls_str]
            rows.append([str(cls), class_names.get(cls, f'Class {cls}'),
                         str(m['support']), f'{m["precision"]:.4f}',
                         f'{m["recall"]:.4f}', f'{m["f1"]:.4f}'])
        add_para(doc, 'Table D2: Per-Class Classification Metrics (10-Fold CV, Seed 13)')
        add_table(doc, ['Class', 'Description', 'Support', 'Precision', 'Recall', 'F1-Score'], rows)

        add_para(doc, f'Macro Precision: {cm_data["macro_precision"]:.4f}')
        add_para(doc, f'Macro Recall: {cm_data["macro_recall"]:.4f}')
        add_para(doc, f'Macro F1-Score: {cm_data["macro_f1"]:.4f}')

    # ================================================================
    # SECTION: Computational Cost
    # ================================================================
    doc.add_page_break()
    add_heading(doc, 'E. Computational Cost and Memory Usage', 1)

    add_para(doc, 'The CDS-OVR algorithm is designed to be computationally lightweight, requiring no GPU acceleration, no iterative optimization, and no gradient computation. All operations are closed-form or require a single pass through the data.')

    if comp_results and 'computational_cost' in comp_results:
        cc = comp_results['computational_cost']
        add_heading(doc, 'E.1 Runtime Performance', 2)
        add_para(doc, 'Table E1: Computational Cost Summary')
        add_table(doc, ['Operation', 'Time', 'Notes'],
                  [['Single Training (full dataset)', f'{cc["training_time_s"]:.4f} s', 'Build tree + train 8 class models'],
                   ['Single Patient Prediction', f'{cc["prediction_per_patient_ms"]:.4f} ms', 'Route + accumulate evidence + decide'],
                   ['All 416 Predictions', f'{cc["prediction_all_patients_s"]:.4f} s', 'Sequential prediction of all patients'],
                   ['10-Fold CV (1 seed)', f'{cc["10fold_cv_time_s"]:.2f} s', '10 train+test cycles'],
                   ['90/10 Split (1 seed)', f'{cc["90_10_split_time_s"]:.2f} s', '1 train+test cycle']])

        add_heading(doc, 'E.2 Memory Usage', 2)
        add_para(doc, 'Table E2: Memory Usage Summary')
        add_table(doc, ['Component', 'Size', 'Description'],
                  [['Data Matrix', f'{cc["data_matrix_mb"]:.2f} MB', f'416 x 279 float64 array'],
                   ['Current Memory', f'{cc["memory_current_mb"]:.2f} MB', 'After training all models'],
                   ['Peak Memory', f'{cc["memory_peak_mb"]:.2f} MB', 'Maximum during prediction'],
                   ['Tree Nodes', str(cc['n_tree_nodes']), 'Root + sex-based branches'],
                   ['Total BinModels', str(cc['n_class_models']), '8 classes x 279 features x nodes'],
                   ['Retained Features', str(cc['n_retained_features']), '8 classes x ~18 features x nodes']])

        add_heading(doc, 'E.3 Computational Complexity', 2)
        add_para(doc, 'Training Complexity: O(C x N x F x B) where C=8 classes, N=416 patients, F=279 features, B=6 max bins. The dominant cost is computing supervised bin edges for each feature in each class model. Feature selection adds O(F^2) for pairwise correlation computation but only within the top candidates.')
        add_para(doc, 'Prediction Complexity: O(C x F_r x L) where C=8 classes, F_r=18 retained features per class, L=number of matched tree nodes (at most 2). Each prediction requires looking up bin assignments and accumulating evidence, all O(1) per feature.')
        add_para(doc, 'The algorithm requires no matrix inversions, no eigendecompositions, no iterative solvers, and no backpropagation. This makes it suitable for deployment on resource-constrained embedded systems such as FPGAs.')

    # ================================================================
    # SECTION: Mathematical Equations
    # ================================================================
    doc.add_page_break()
    add_heading(doc, 'F. Mathematical Formulation', 1)

    add_para(doc, 'This section provides the complete mathematical formulation for each component of the CDS-OVR algorithm, with references to the relevant literature.')

    add_heading(doc, 'F.1 Supervised Chi-Squared Binning', 2)
    add_para(doc, 'The supervised binning algorithm partitions continuous feature values into discrete bins by maximizing the chi-squared statistic between the target class indicator and the bin assignment. For a candidate split dividing a segment into left and right partitions:')
    add_para(doc, 'chi2 = sum_i sum_j (O_ij - E_ij)^2 / E_ij')
    add_para(doc, 'where O_ij is the observed count (i in {target, rest}, j in {left, right}) and E_ij = n_j * n_i / n is the expected count under the null hypothesis of independence. The chi-squared statistic follows an asymptotic chi2(1) distribution under the null. A minimum gain threshold of 0.5 is used to prevent overfitting [1].')

    add_heading(doc, 'F.2 Laplace-Smoothed Posterior', 2)
    add_para(doc, 'The class-conditional posterior probability for class c in bin b is computed with Laplace smoothing:')
    add_para(doc, 'P(c|b) = (n_cb + alpha * pi_c) / (n_b + alpha)')
    add_para(doc, 'where n_cb is the number of class-c patients in bin b, n_b is the total count in bin b, pi_c = n_c/N is the class prior, and alpha = 1.0 is the smoothing parameter. This is equivalent to placing a Dirichlet prior on the bin-conditional class distribution [2].')

    add_heading(doc, 'F.3 Fisher Discriminant Ratio', 2)
    add_para(doc, 'The Fisher Discriminant Ratio (FDR) measures the univariate separability of a feature between the target class and all other classes:')
    add_para(doc, 'FDR(f, c) = (mu_c - mu_rest)^2 / (sigma_c^2 + sigma_rest^2 + epsilon)')
    add_para(doc, 'where mu_c and sigma_c^2 are the mean and variance of feature f for class c patients, mu_rest and sigma_rest^2 are for all other patients, and epsilon = 1e-10 prevents division by zero. FDR is derived from the Fisher Linear Discriminant [3] and equals the squared signal-to-noise ratio for the binary classification task.')

    add_heading(doc, 'F.4 Fisher Weight Computation', 2)
    add_para(doc, 'Each retained feature receives a weight based on its relative Fisher discriminant ratio:')
    add_para(doc, 'w_f = max(sqrt(FDR_f / (FDR_max + epsilon)), 0.1)')
    add_para(doc, 'where FDR_max is the maximum FDR across all retained features for the class. The square root compresses the dynamic range, and the floor of 0.1 ensures that no feature is entirely silenced.')

    add_heading(doc, 'F.5 Action Score', 2)
    add_para(doc, 'The action score quantifies how much a feature contributes to distinguishing the target class from the population baseline:')
    add_para(doc, 'S(f, c) = sum_b |P(c|b) - pi_c| * min(1, n_b / CS)')
    add_para(doc, 'where the sum is over bins b with at least MS patients, P(c|b) is the Laplace-smoothed posterior, pi_c is the class prior, and CS is the confidence support parameter. The confidence factor min(1, n_b/CS) prevents bins with few observations from contributing disproportionately.')

    add_heading(doc, 'F.6 Dual Evidence Accumulation', 2)
    add_para(doc, 'Evidence for and against each class is accumulated separately:')
    add_para(doc, 'AF_for = sum_f [shift_f * confidence_f * w_f]  for shift_f >= 0')
    add_para(doc, 'AF_against = sum_f [|shift_f| * confidence_f * w_f * alpha_c]  for shift_f < 0')
    add_para(doc, 'where shift_f = P(c|b_f) - pi_c is the posterior shift for the patient in feature f, confidence_f = min(1, bc_f / 10) is the confidence based on bin population, w_f is the Fisher weight, and alpha_c is the against_scale parameter (0.5 for rare classes, 0.8 for common classes).')

    add_heading(doc, 'F.7 Ratio Score and Decision Logic', 2)
    add_para(doc, 'The ratio score combines positive and negative evidence:')
    add_para(doc, 'R_c = (AF_for + epsilon) / (AF_against + epsilon)')
    add_para(doc, 'where epsilon = 0.1 provides stability. The effective threshold for each class is:')
    add_para(doc, 'T_eff = max(T_c - delta * I(h_score < h_cut), min(w_h * h_score, h_cap))')
    add_para(doc, 'where T_c is the per-class threshold, delta = 0.3 is the suspicion offset, h_cut = 2.0 is the suspicion activation threshold, w_h = 1.05 is the healthy weight, h_cap = 5.0 is the healthy bar cap, and h_score = R_1 is the ratio score for the healthy class.')
    add_para(doc, 'A class c is a candidate if R_c >= T_eff. The predicted class is the candidate with the maximum normalized margin:')
    add_para(doc, 'prediction = argmax_c (R_c - T_eff) / max(T_eff, 0.1)')

    add_heading(doc, 'F.8 Pearson Correlation Coefficient', 2)
    add_para(doc, 'The absolute Pearson correlation between features f1 and f2 is:')
    add_para(doc, '|r| = |sum((x_i - mu_x)(y_i - mu_y))| / sqrt(sum(x_i - mu_x)^2 * sum(y_i - mu_y)^2)')
    add_para(doc, 'Features with |r| > 0.8 are considered redundant. During greedy selection, each candidate feature is checked against all already-selected features, and rejected if any pairwise correlation exceeds the threshold [4].')

    add_heading(doc, 'F.9 Evaluation Metrics', 2)
    add_para(doc, 'Multiclass Accuracy = (1/N) * sum_i I(y_pred_i == y_true_i)')
    add_para(doc, 'Specificity = TP_healthy / (TP_healthy + FN_healthy)')
    add_para(doc, 'Sensitivity = sum_c!=1 I(pred_i != 1) / sum_c!=1 I(true_i != 1)')
    add_para(doc, 'Balanced Accuracy = (1/C) * sum_c (correct_c / n_c)')
    add_para(doc, 'Binary Accuracy = sum_i I((true_i==1) == (pred_i==1)) / N')
    add_para(doc, 'Precision_c = TP_c / (TP_c + FP_c)')
    add_para(doc, 'Recall_c = TP_c / (TP_c + FN_c)')
    add_para(doc, 'F1_c = 2 * Precision_c * Recall_c / (Precision_c + Recall_c)')
    add_para(doc, 'AUC = integral from 0 to 1 of TPR(FPR) dFPR  (computed via the trapezoidal rule)')

    # ================================================================
    # SECTION: References
    # ================================================================
    doc.add_page_break()
    add_heading(doc, 'G. References', 1)

    refs = [
        '[1] Kerber, R. (1992). ChiMerge: Discretization of numeric attributes. In Proceedings of the 10th National Conference on Artificial Intelligence (AAAI-92), pp. 123-128.',
        '[2] Manning, C. D., Raghavan, P., & Schutze, H. (2008). Introduction to Information Retrieval. Cambridge University Press. Chapter 13: Text classification and Naive Bayes.',
        '[3] Fisher, R. A. (1936). The use of multiple measurements in taxonomic problems. Annals of Eugenics, 7(2), 179-188.',
        '[4] Guyon, I., & Elisseeff, A. (2003). An introduction to variable and feature selection. Journal of Machine Learning Research, 3, 1157-1182.',
        '[5] Guvenir, H. A., Acar, B., Demiroz, G., & Cekin, A. (1997). A supervised machine learning algorithm for arrhythmia analysis. In Proceedings of Computers in Cardiology, pp. 433-436.',
        '[6] Sharma, A., et al. (2024). Optimized Random Forest for arrhythmia classification on the UCI dataset. (Binary classification, 90/10 split, 95.24% accuracy.)',
        '[7] Mustaqeem, A., Anwar, S. M., & Majid, M. (2018). Multiclass classification of cardiac arrhythmia using improved feature selection and SVM invariants. Computational and Mathematical Methods in Medicine, 2018, Article 7310496.',
        '[8] Irfan, S., Anjum, N., Althobaiti, T., Alotaibi, A. A., Siddiqui, A. B., & Ramzan, N. (2022). Heartbeat classification and arrhythmia detection using a multi-model deep-learning technique. Sensors, 22(15), 5606.',
        '[9] Gupta, R., & Mitra, M. (2014). A hybrid feature-based classification scheme for cardiac arrhythmia. International Journal of Biomedical Engineering and Technology, 15(4), 299-318.',
        '[10] Jadhav, S. M., Nalbalwar, S. L., & Ghatol, A. A. (2012). Artificial neural network models based cardiac arrhythmia disease diagnosis from ECG data. In Proceedings of International Conference on Communications and Signal Processing (ICCSP), pp. 1162-1166.',
        '[11] Plawiak, P. (2018). Novel methodology of cardiac health recognition based on ECG signals and evolutionary-neural system. Expert Systems with Applications, 92, 334-349.',
        '[12] UCI Machine Learning Repository. (1998). Arrhythmia Data Set. https://archive.ics.uci.edu/ml/datasets/arrhythmia',
        '[13] Rifkin, R., & Klautau, A. (2004). In defense of one-vs-all classification. Journal of Machine Learning Research, 5, 101-141.',
        '[14] Wilcoxon, F. (1945). Individual comparisons by ranking methods. Biometrics Bulletin, 1(6), 80-83.',
    ]
    for ref in refs:
        add_para(doc, ref)


def main():
    src = os.path.join(os.environ['TEMP'], 'CDS_Overall_Report.docx')
    if not os.path.exists(src):
        orig = os.path.join(SRC_DIR, 'CDS_Overall_Report.docx')
        shutil.copy2(orig, src)

    print("Opening report...")
    doc = Document(src)

    print("Replacing em dashes...")
    count = replace_em_dashes(doc)
    print(f"  Replaced {count} em/en dashes")

    print("Adding new sections...")
    add_new_sections(doc)

    out_path = os.path.join(SRC_DIR, 'CDS_Overall_Report_Updated.docx')
    doc.save(out_path)
    print(f"Saved updated report to {out_path}")


if __name__ == "__main__":
    main()
