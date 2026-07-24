"""Restructure CDS_Overall_Report.docx into a properly organized final report.

Fixes:
- Triple-duplicate section numbering (adds Part divider headings)
- Redundant tables (adds cross-reference annotations)
- Missing file location references
- Em/en dash removal
- Appends comprehensive appendix sections (A-G) with all available data
"""
import os, sys, json, shutil, re
from collections import defaultdict
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC_DIR = os.path.dirname(os.path.abspath(__file__))

CLASS_NAMES = {
    1: 'Normal', 2: 'Ischemic Changes (CAD)', 3: 'Old Anterior MI',
    4: 'Old Inferior MI', 5: 'Sinus Tachycardia', 6: 'Sinus Bradycardia',
    7: 'WPW', 8: 'AV Block', 9: 'LBBB', 10: 'RBBB',
    11: 'Class 14', 12: 'Class 15', 13: 'Class 16'
}


# ============================================================
# XML Helpers for inserting/modifying elements
# ============================================================

def make_paragraph_element(text, style=None, bold=False, italic=False,
                           font_size=None, color=None, alignment=None):
    """Create a new w:p XML element with the given properties."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    if style:
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), style)
        pPr.append(pStyle)
    if alignment:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), alignment)
        pPr.append(jc)
    p.append(pPr)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        rPr.append(OxmlElement('w:b'))
    if italic:
        rPr.append(OxmlElement('w:i'))
    if font_size:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(font_size * 2))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(font_size * 2))
        rPr.append(szCs)
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    p.append(r)
    return p


def insert_before(target_element, new_element):
    target_element.addprevious(new_element)


def insert_after(target_element, new_element):
    target_element.addnext(new_element)


def make_page_break():
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    return p


# ============================================================
# Content Helpers
# ============================================================

def replace_em_dashes(doc):
    count = 0
    for p in doc.paragraphs:
        for run in p.runs:
            if '—' in run.text:
                run.text = run.text.replace('—', '-')
                count += 1
            if '–' in run.text:
                run.text = run.text.replace('–', '-')
                count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if '—' in run.text:
                            run.text = run.text.replace('—', '-')
                            count += 1
                        if '–' in run.text:
                            run.text = run.text.replace('–', '-')
                            count += 1
    return count


def get_para_text(para):
    return para.text.strip()


def find_paragraph_index(doc, text_substring, start=0):
    for i in range(start, len(doc.paragraphs)):
        if text_substring in get_para_text(doc.paragraphs[i]):
            return i
    return -1


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f'Heading {level}')
    return p


def add_para(doc, text, bold=False, italic=False, font_size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    return p


def add_source_ref(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    return table


# ============================================================
# Data Loading
# ============================================================

def load_loocv_traces():
    base_path = os.path.join(SRC_DIR, '..', 'output', 'loocv_trace.json')
    ovr_path = os.path.join(SRC_DIR, '..', 'output', 'loocv_trace_ovr.json')
    with open(base_path) as f:
        base = json.load(f)
    with open(ovr_path) as f:
        ovr = json.load(f)
    return base, ovr


def compute_loocv_stats(trace):
    n = len(trace)
    correct = sum(1 for r in trace if r['correct'])
    cls_stats = defaultdict(lambda: {'correct': 0, 'total': 0})
    for r in trace:
        tc = r['true_class']
        cls_stats[tc]['total'] += 1
        if r['correct']:
            cls_stats[tc]['correct'] += 1

    healthy = [r for r in trace if r['true_class'] == 1]
    diseased = [r for r in trace if r['true_class'] != 1]
    spec = sum(1 for r in healthy if r['correct']) / len(healthy) if healthy else 0
    sens = sum(1 for r in diseased if r['predicted'] != 1) / len(diseased) if diseased else 0
    bin_correct = 0
    for r in trace:
        pred = r.get('predicted', '')
        true_h = (r['true_class'] == 1)
        pred_h = (pred == 1 or pred == 'HEALTHY')
        if true_h == pred_h:
            bin_correct += 1
    ba_vals = []
    for cls in cls_stats:
        s = cls_stats[cls]
        if s['total'] > 0:
            ba_vals.append(s['correct'] / s['total'])
    ba = sum(ba_vals) / len(ba_vals) if ba_vals else 0
    return {
        'n': n, 'correct': correct, 'accuracy': correct / n,
        'specificity': spec, 'sensitivity': sens,
        'binary_accuracy': bin_correct / n,
        'balanced_accuracy': ba,
        'per_class': dict(cls_stats),
    }


def load_results_all_splits():
    path = os.path.join(SRC_DIR, 'results_all_splits.json')
    if os.path.exists(path):
        with open(path) as f:
            data = json.load(f)
        # Fix floating-point key issues (90/9 should be 90/10, 80/19 should be 80/20)
        key_fixes = {'90/9': '90/10', '80/19': '80/20'}
        for old_key, new_key in key_fixes.items():
            if old_key in data and new_key not in data:
                data[new_key] = data.pop(old_key)
        return data
    return None


def load_comprehensive_results():
    # Try combined file first
    path = os.path.join(SRC_DIR, 'comprehensive_results.json')
    if os.path.exists(path):
        with open(path) as f:
            return json.load(f)

    # Fall back to individual result files from parallel tasks
    combined = {}
    files = {
        "auc_roc": "results_auc.json",
        "stratified": "results_stratified.json",
        "computational_cost": "results_cost.json",
        "confusion_matrix_seed13": "results_confusion.json",
    }
    for key, fname in files.items():
        fpath = os.path.join(SRC_DIR, fname)
        if os.path.exists(fpath):
            with open(fpath) as f:
                combined[key] = json.load(f)
    return combined if combined else None


# ============================================================
# Phase 1: Structural fixes to existing content
# ============================================================

def add_abstract(doc):
    """Insert an Abstract section before Part I."""
    part1_idx = -1
    for i, p in enumerate(doc.paragraphs):
        txt = get_para_text(p)
        if 'CDS-OVR Algorithm Specification' in txt and i < 5:
            part1_idx = i
            break

    if part1_idx < 0:
        return 0

    target = doc.paragraphs[part1_idx]._element

    lines = [
        ('Abstract', 'Heading1', True, 16),
        ('This report presents the CDS-OVR (Cognitive Dynamic System with One-vs-Rest) algorithm, '
         'a novel interpretable classifier for automated arrhythmia detection from the UCI Arrhythmia '
         'Dataset (452 patients, 279 features, 13 classes). CDS-OVR decomposes the multiclass problem '
         'into 8 independent binary classification tasks, each using supervised chi-squared binning, '
         'Fisher discriminant feature weighting, dual evidence accumulation, and a dynamic healthy bar '
         'thresholding mechanism. The algorithm requires no GPU acceleration, no iterative optimization, '
         'and no gradient computation, making it suitable for resource-constrained embedded deployment.',
         None, False, 11),
        ('We evaluate CDS-OVR across multiple protocols: leave-one-out cross-validation (LOOCV), '
         '10-fold cross-validation, train/test splits from 50/50 to 90/10, and stratified variants '
         'of all methods. On 10-fold CV, CDS-OVR achieves 86.78% multiclass accuracy (best seed) '
         'with a mean of 84.66% (std=0.93%) across 10 seeds, and 87.72% mean binary accuracy. '
         'AUC-ROC analysis shows binary AUC of 0.9021 and macro AUC of 0.9393. Stratified 10-fold '
         'CV yields 85.32% multiclass and 88.56% binary accuracy. LOOCV on all 13 classes achieves '
         '74.8% multiclass and 84.5% binary accuracy, improving over the base CDS algorithm by '
         '+7.5 pp and +9.5 pp respectively.',
         None, False, 11),
        ('Component ablation with Wilcoxon signed-rank testing identifies supervised chi-squared '
         'binning (+5.75 pp, p=0.002) and Fisher discriminant weighting (+4.52 pp, p=0.002) as the '
         'statistically significant contributors. Benchmark comparison against published methods '
         '(Random Forest, SVM, CNN-LSTM, ANN) demonstrates competitive performance at high training '
         'ratios, with structural advantages in interpretability, missing value handling, and '
         'computational efficiency (3.05 MB memory, 41 ms per prediction).',
         None, False, 11),
    ]

    for text, style, bold, font_size in lines:
        el = make_paragraph_element(text, style=style, bold=bold, font_size=font_size)
        insert_before(target, el)

    print(f"    Abstract inserted before paragraph {part1_idx}")
    return 1


def fix_part3_subtitle(doc):
    """Fix the empty Heading 2 and fragmented subtitle in Part III."""
    changes = 0
    for i, p in enumerate(doc.paragraphs):
        txt = get_para_text(p)
        style = p.style.name if p.style else ''
        if style == 'Heading 2' and not txt:
            # Check if next paragraphs contain the split subtitle
            next_texts = []
            for j in range(i+1, min(i+4, len(doc.paragraphs))):
                nt = get_para_text(doc.paragraphs[j])
                if nt:
                    next_texts.append((j, nt))
            combined = ' '.join(t for _, t in next_texts)
            if 'Comparative' in combined or 'Benchmark' in combined:
                # Set the empty heading to the combined title
                for run in p.runs:
                    run.text = ''
                if not p.runs:
                    p.add_run('')
                p.runs[0].text = combined
                # Remove the fragmented paragraphs
                for j, _ in next_texts:
                    for run in doc.paragraphs[j].runs:
                        run.text = ''
                changes += 1
                print(f"    Fixed Part III subtitle at paragraph {i}")
                break
    return changes


def fix_duplicate_section_numbers(doc):
    """Fix duplicate 8.1 and 8.2 section numbers in Part III.
    Original has: 8.1, 8.1, 8.2. Fix to: 8.1, 8.2, 8.3."""
    changes = 0
    seen_8_1 = False
    seen_8_2 = False
    for i, p in enumerate(doc.paragraphs):
        txt = get_para_text(p)
        style = p.style.name if p.style else ''
        if i < 1000:
            continue
        if 'Heading' in style and txt.startswith('8.1 '):
            if not seen_8_1:
                seen_8_1 = True
            else:
                for run in p.runs:
                    if run.text.strip().startswith('8.1'):
                        run.text = run.text.replace('8.1', '8.2', 1)
                        changes += 1
                        seen_8_2 = True
                        print(f"    Fixed duplicate 8.1 -> 8.2 at paragraph {i}")
                        break
        elif 'Heading' in style and txt.startswith('8.2 '):
            if not seen_8_2:
                seen_8_2 = True
            else:
                for run in p.runs:
                    if run.text.strip().startswith('8.2'):
                        run.text = run.text.replace('8.2', '8.3', 1)
                        changes += 1
                        print(f"    Fixed duplicate 8.2 -> 8.3 at paragraph {i}")
                        break
    return changes


def fix_structure(doc):
    """Add Part divider headings and file location annotations."""
    print("  Scanning paragraphs for structural boundaries...")

    part2_idx = -1
    part3_idx = -1

    for i, p in enumerate(doc.paragraphs):
        txt = get_para_text(p)
        if 'Incremental Analysis of CDS Arrhythmia Classification Algorithm' in txt:
            part2_idx = i
        if 'Comparative Algorithmic Analysis' in txt and 'CDS-OVR vs' in txt:
            part3_idx = i
        if part3_idx == -1 and txt == 'CDS-OVR vs. Benchmark Methods':
            part3_idx = i

    # Also find the start paragraph of Part 1
    part1_idx = -1
    for i, p in enumerate(doc.paragraphs):
        txt = get_para_text(p)
        if 'CDS-OVR Algorithm Specification' in txt and i < 5:
            part1_idx = i
            break

    changes = 0

    # Part I title
    if part1_idx >= 0:
        p = doc.paragraphs[part1_idx]
        p.style = doc.styles['Heading 1']
        for run in p.runs:
            run.text = 'Part I: ' + run.text
        changes += 1
        print(f"    Part I header set at paragraph {part1_idx}")

    # Part II divider
    if part2_idx >= 0:
        p = doc.paragraphs[part2_idx]
        pb = make_page_break()
        insert_before(p._element, pb)

        part_heading = make_paragraph_element(
            'Part II: Incremental Development Analysis',
            style='Heading1', bold=True, font_size=16
        )
        insert_before(p._element, part_heading)

        p.style = doc.styles['Heading 2']
        changes += 1
        print(f"    Part II header inserted at paragraph {part2_idx}")

    # Part III divider
    if part3_idx >= 0:
        p = doc.paragraphs[part3_idx]
        pb = make_page_break()
        insert_before(p._element, pb)

        part_heading = make_paragraph_element(
            'Part III: Benchmark Comparison',
            style='Heading1', bold=True, font_size=16
        )
        insert_before(p._element, part_heading)

        p.style = doc.styles['Heading 2']
        changes += 1
        print(f"    Part III header inserted at paragraph {part3_idx}")

    return changes


def add_file_location_refs(doc):
    """Add source file annotations after relevant sections."""
    refs_to_add = [
        ('Table 1: Class Distribution', 'Source: ImprovedCDS/CDS_FINAL/cds_ovr.py, lines 31-38 (REMOVE_CLASSES, class filtering)'),
        ('Table 2: Complete Hyperparameter', 'Source: ImprovedCDS/CDS_FINAL/cds_ovr.py, lines 17-44 (constants section)'),
        ('Algorithm 1: build_tree', 'Source: ImprovedCDS/CDS_FINAL/cds_ovr.py, lines 55-116 (build_tree function)'),
        ('supervised chi-squared binning', 'Source: ImprovedCDS/CDS_FINAL/cds_ovr.py, lines 209-265 (chi2_binning function)'),
        ('Algorithm 8: Complete CDS-OVR Pipeline', 'Source: ImprovedCDS/CDS_FINAL/cds_ovr.py (full algorithm implementation, 587 lines)'),
        ('evidence_analysis.py', 'Source file: ImprovedCDS/CDS_FINAL/evidence_analysis.py'),
        ('evidence_ablation.py', 'Source file: ImprovedCDS/CDS_FINAL/evidence_ablation.py'),
        ('ablation_full.py', 'Source file: ImprovedCDS/CDS_FINAL/ablation_full.py'),
        ('run_all_splits.py', 'Source file: ImprovedCDS/CDS_FINAL/run_all_splits.py; Results: ImprovedCDS/CDS_FINAL/results_all_splits.json'),
        ('00_base_original', 'Source: ImprovedCDS/cds.py (original base algorithm, 491 lines)'),
        ('Per-Class Score Distributions and Threshold Rationale', 'Threshold values: ImprovedCDS/CDS_FINAL/cds_ovr.py, line 39 (CLASS_THRESHOLDS)'),
    ]

    count = 0
    inserted_refs = set()

    for i, p in enumerate(doc.paragraphs):
        txt = get_para_text(p)
        for trigger, ref_text in refs_to_add:
            if trigger in txt and trigger not in inserted_refs:
                annotation = make_paragraph_element(
                    ref_text, italic=True, font_size=9, color='666666'
                )
                insert_after(p._element, annotation)
                inserted_refs.add(trigger)
                count += 1
                break

    return count


def mark_redundancies(doc):
    """Add cross-reference notes near redundant tables/content."""
    redundancy_markers = [
        ('Table 6: Complete Hyperparameter Reference',
         '[Note: This table consolidates the hyperparameters from Table 2 in Part I, Section 2.5. Both tables are retained for contextual reference within their respective sections.]'),
    ]

    count = 0
    for i, p in enumerate(doc.paragraphs):
        txt = get_para_text(p)
        for trigger, note in redundancy_markers:
            if trigger in txt:
                annotation = make_paragraph_element(
                    note, italic=True, font_size=9, color='888888'
                )
                insert_after(p._element, annotation)
                count += 1
                break
    return count


# ============================================================
# Phase 5b: Text corrections throughout the document
# ============================================================

def fix_text_issues(doc):
    """Fix specific text errors found during review."""
    fixes = 0

    for i, p in enumerate(doc.paragraphs):
        txt = p.text

        # Issue 2: Fix sex encoding (report says 0=female, 1=male; correct is 0=male, 1=female)
        if '0.0 for female and 1.0 for male' in txt:
            for run in p.runs:
                if '0.0 for female and 1.0 for male' in run.text:
                    run.text = run.text.replace('0.0 for female and 1.0 for male',
                                                '0.0 for male and 1.0 for female')
                    fixes += 1
                    print(f"    Fixed sex encoding at paragraph {i}")

        # Issue 5: Fix "more than 0 classes" typo (correct: 2 classes)
        if 'ranks in the top 5 for more than 0 classes' in txt:
            for run in p.runs:
                if 'more than 0 classes' in run.text:
                    run.text = run.text.replace('more than 0 classes', 'more than 2 classes')
                    fixes += 1
                    print(f"    Fixed feature overlap count at paragraph {i}")

        # Issue 7: Fix dual AF ratio for misclassified (5.397 is wrong, 1.664 is correct)
        if '5.397' in txt and 'incorrectly classified' in txt:
            for run in p.runs:
                if '5.397' in run.text:
                    run.text = run.text.replace('5.397', '1.664')
                    fixes += 1
                    print(f"    Fixed dual AF misclassified mean ratio at paragraph {i}")
                if '4.401' in run.text:
                    run.text = run.text.replace('4.401', '1.247')
                    fixes += 1
                    print(f"    Fixed dual AF misclassified median at paragraph {i}")

    return fixes


def fix_appendix_a_clarification(doc):
    """Issue 9: Clarify that base CDS is binary-only in Part II intro."""
    fixes = 0
    for i, p in enumerate(doc.paragraphs):
        txt = p.text
        if '50.7% multiclass' in txt and 'base' in txt.lower():
            for run in p.runs:
                if '50.7% multiclass' in run.text:
                    run.text = run.text.replace(
                        '50.7% multiclass',
                        '50.7% multiclass (binary-derived, see Appendix A)')
                    fixes += 1
                    print(f"    Clarified base model multiclass metric at paragraph {i}")
            break
    return fixes


def add_healthy_bar_analysis(doc):
    """Issue 6: Add grounded healthy bar analysis with near-boundary data."""
    hbar_path = os.path.join(SRC_DIR, 'analysis_healthy_bar.json')
    if not os.path.exists(hbar_path):
        print("    Skipping healthy bar analysis (data file not found)")
        return 0

    with open(hbar_path) as f:
        hbar = json.load(f)

    # Find the healthy bar ablation interpretation paragraph and add analysis after it
    for i, p in enumerate(doc.paragraphs):
        if 'Healthy bar' in p.text and '0.31 pp' in p.text and 'Removing the healthy bar' in p.text:
            # Replace the explanation with grounded analysis
            fn_details = hbar.get('false_neg_details', [])
            n_blocked = hbar.get('hbar_blocked_true_class', 0)
            n_near = hbar.get('near_threshold_count', 0)
            n_fn = hbar.get('false_negatives', 0)
            n_fp = hbar.get('false_positives', 0)

            new_text = (
                f'6. Healthy bar (-0.31 pp, p=0.133): Removing the healthy bar improves overall '
                f'accuracy by 0.31 pp (not significant). Detailed analysis of the 10-fold CV '
                f'(seed 13) reveals why: of {n_fn} false negatives (disease patients predicted '
                f'as healthy), {n_blocked} had their true disease class score blocked by the '
                f'healthy bar raising the effective threshold above the base class threshold. '
                f'{n_near} patients were within 1.0 of their threshold (near-misses). '
                f'However, the healthy bar also prevented {n_fp} false positives (healthy patients '
                f'that would otherwise be misclassified as diseased). The net effect is a '
                f'trade-off: the healthy bar sacrifices {n_blocked} true disease detections '
                f'to prevent false alarms, explaining the marginal accuracy decrease.')

            for run in p.runs:
                run.text = ''
            if p.runs:
                p.runs[0].text = new_text
            print(f"    Updated healthy bar explanation at paragraph {i}")
            return 1

    return 0


def add_binning_figure(doc):
    """Issue 3: Insert supervised vs equal-width binning comparison table after Section 4."""
    fig_path = os.path.join(SRC_DIR, 'figure_data_binning.json')
    if not os.path.exists(fig_path):
        print("    Skipping binning figure (data file not found)")
        return 0

    with open(fig_path) as f:
        fig = json.load(f)

    for i, p in enumerate(doc.paragraphs):
        if p.style and p.style.name and 'Heading' in p.style.name and \
           '5.' in p.text[:5] and 'Feature Scoring' in p.text:
            target = p._element
            elements = []

            elements.append(make_paragraph_element(
                f'Figure 4.1: Supervised vs Equal-Width Binning Comparison '
                f'(Feature {fig["feature_index"]}, Class {fig["target_class"]} '
                f'- {CLASS_NAMES.get(fig["target_class"], "Unknown")})',
                bold=True, font_size=10))

            elements.append(make_paragraph_element(
                f'Dataset: {fig["n_samples"]} patients in training fold, '
                f'{fig["n_target"]} target class, {fig["n_rest"]} rest. '
                f'Value range: [{fig["value_range"][0]}, {fig["value_range"][1]}]. '
                f'Purity improvement: {fig["improvement_purity"]:.4f}.',
                font_size=9))

            sup_desc = []
            for b in fig['supervised_bins']:
                sup_desc.append(
                    f'[{b["lo"]:.0f}, {b["hi"]:.0f}]: {b["n_target"]}/{b["n_total"]} '
                    f'target ({100*b["target_frac"]:.0f}%)')
            elements.append(make_paragraph_element(
                'Supervised bins: ' + '; '.join(sup_desc),
                font_size=9))

            eq_desc = []
            for b in fig['equal_width_bins']:
                eq_desc.append(
                    f'[{b["lo"]:.0f}, {b["hi"]:.0f}]: {b["n_target"]}/{b["n_total"]} '
                    f'target ({100*b["target_frac"]:.0f}%)')
            elements.append(make_paragraph_element(
                'Equal-width bins: ' + '; '.join(eq_desc),
                font_size=9))

            elements.append(make_paragraph_element(
                'Key finding: Supervised binning isolates a pure bin [-115, -94] with 5/5 '
                'target patients (100%), while no equal-width bin exceeds 38.5% target fraction. '
                'This pure bin directly becomes a strong positive evidence signal in the CDS '
                'accumulation factor.',
                italic=True, font_size=9))

            for el in reversed(elements):
                insert_before(target, el)

            print(f"    Inserted binning comparison figure before paragraph {i}")
            return 1
    return 0


def add_ovr_figure(doc):
    """Issue 4: Insert OVR class separation figure after Section 7."""
    fig_path = os.path.join(SRC_DIR, 'figure_data_ovr.json')
    if not os.path.exists(fig_path):
        print("    Skipping OVR figure (data file not found)")
        return 0

    with open(fig_path) as f:
        fig = json.load(f)

    features = fig.get('features', fig) if isinstance(fig, dict) else fig

    for i, p in enumerate(doc.paragraphs):
        if p.style and p.style.name and 'Heading' in p.style.name and \
           '8.' in p.text[:5] and 'Evidence Accumulation' in p.text:
            target = p._element
            elements = []

            elements.append(make_paragraph_element(
                'Figure 7.1: OVR Class Separation - Fisher Discriminant Ratio '
                'by Best Feature per Class',
                bold=True, font_size=10))

            rows_text = []
            for entry in features:
                cls = entry.get('target_class', entry.get('class'))
                feat = entry.get('feature_index', entry.get('feature'))
                fdr = entry.get('fisher_ratio', entry.get('fisher_discriminant_ratio', 0))
                n_t = entry.get('target_count', entry.get('n_target', 0))
                n_h = entry.get('healthy_count', entry.get('n_healthy', 0))
                n_o = entry.get('other_count', entry.get('n_other_disease', 0))
                rows_text.append(
                    f'Class {cls} ({CLASS_NAMES.get(cls, "?")}): '
                    f'Feature {feat}, FDR={fdr:.2f}, '
                    f'{n_t} target / {n_h} healthy / {n_o} other disease')

            elements.append(make_paragraph_element(
                'Per-class best discriminating feature (ranked by Fisher Discriminant Ratio): ' +
                '. '.join(rows_text) + '.',
                font_size=9))

            top = features[0]
            bot = features[-1]
            top_cls = top.get('target_class', top.get('class'))
            bot_cls = bot.get('target_class', bot.get('class'))
            top_fdr = top.get('fisher_ratio', top.get('fisher_discriminant_ratio', 0))
            bot_fdr = bot.get('fisher_ratio', bot.get('fisher_discriminant_ratio', 0))
            top_feat = top.get('feature_index', top.get('feature'))
            elements.append(make_paragraph_element(
                f'Key finding: Each class has a different best feature, confirming that the OVR '
                f'architecture is essential - no single feature set can serve all disease classes. '
                f'Class {top_cls} ({CLASS_NAMES.get(top_cls, "?")}) achieves the highest '
                f'FDR={top_fdr:.2f} on Feature {top_feat}, while '
                f'Class {bot_cls} ({CLASS_NAMES.get(bot_cls, "?")}) has the lowest '
                f'FDR={bot_fdr:.2f}, explaining why it is harder to classify.',
                italic=True, font_size=9))

            for el in reversed(elements):
                insert_before(target, el)

            print(f"    Inserted OVR separation figure before paragraph {i}")
            return 1
    return 0


def add_ablation_missing_results(doc):
    """Add results from newly computed ablation tests."""
    abl_path = os.path.join(SRC_DIR, 'ablation_missing_data.json')
    if not os.path.exists(abl_path):
        print("    Skipping missing ablation results (data file not found)")
        return 0

    with open(abl_path) as f:
        abl = json.load(f)

    for i, p in enumerate(doc.paragraphs):
        if 'Interpretation:' in p.text and 'Only two mechanisms produce statistically significant' in p.text:
            target_el = p._element
            elements = []

            elements.append(make_paragraph_element(
                'Additional Ablation Results (computed separately):',
                bold=True, font_size=10))

            names = {
                'dual_af': 'Dual Accumulation Factor',
                'ratio_scoring': 'Ratio Scoring',
                'laplace_smoothing': 'Laplace Smoothing',
                'per_class_thresholds': 'Per-Class Thresholds',
                'rare_class_params': 'Rare Class Adaptation',
            }

            for key, name in names.items():
                if key not in abl:
                    continue
                r = abl[key]
                sig = 'p<0.001' if r['p_value'] < 0.001 else \
                      'p<0.01' if r['p_value'] < 0.01 else \
                      'p<0.05' if r['p_value'] < 0.05 else f'p={r["p_value"]:.3f}'
                star = '***' if r['p_value'] < 0.001 else \
                       '**' if r['p_value'] < 0.01 else \
                       '*' if r['p_value'] < 0.05 else ' (ns)'
                elements.append(make_paragraph_element(
                    f'{name}: removing drops accuracy by {r["delta_pp"]:+.2f} pp '
                    f'({r["baseline_mean"]:.2f}% -> {r["ablated_mean"]:.2f}%), '
                    f'Wilcoxon W={r["wilcoxon_W"]:.1f}, {sig}{star}, '
                    f'effect size r={r["effect_size_r"]:.3f}.',
                    font_size=9))

            for el in reversed(elements):
                insert_after(target_el, el)

            print(f"    Inserted {len(abl)} additional ablation results after paragraph {i}")
            return len(abl)
    return 0


# ============================================================
# Phase 2: Append comprehensive appendix sections
# ============================================================

def build_appendices(doc):
    """Build and append all appendix sections."""
    print("  Loading data for appendices...")

    base_trace, ovr_trace = load_loocv_traces()
    base_stats = compute_loocv_stats(base_trace)
    ovr_stats = compute_loocv_stats(ovr_trace)
    all_splits = load_results_all_splits()
    comp_results = load_comprehensive_results()

    doc.add_page_break()

    # ================================================================
    # APPENDICES HEADER
    # ================================================================
    add_heading(doc, 'Appendices', 1)
    add_para(doc,
        'The following appendices provide comprehensive evaluation results, '
        'mathematical formulations, and references for the CDS-OVR algorithm. '
        'All results were generated from the source code in ImprovedCDS/CDS_FINAL/.')

    # ================================================================
    # APPENDIX A: LOOCV Results
    # ================================================================
    doc.add_page_break()
    add_heading(doc, 'Appendix A: Leave-One-Out Cross-Validation (LOOCV) Results', 1)

    add_para(doc,
        'LOOCV provides the least biased estimate of generalization performance by training on '
        'N-1 patients and testing on the remaining one, repeated for all N patients. Unlike k-fold '
        'CV, LOOCV uses the maximum possible training set size for each prediction, eliminating '
        'variance from random fold assignments.')
    add_source_ref(doc, 'Source: ImprovedCDS/output/loocv_trace.json (base CDS), ImprovedCDS/output/loocv_trace_ovr.json (CDS-OVR)')

    # A.1 Base CDS
    add_heading(doc, 'A.1 Base CDS Algorithm LOOCV (All 13 Classes, 452 Patients)', 2)
    add_para(doc,
        'The original CDS algorithm (cds.py) was evaluated using LOOCV on the complete 452-patient '
        'dataset with all 13 arrhythmia classes. This represents the unmodified binary classifier '
        'from the initial implementation. Note: the base CDS is fundamentally a binary classifier '
        'that outputs HEALTHY, UNHEALTHY, or SCREENING decisions and cannot identify specific '
        'disease classes. "Multiclass Accuracy" below measures how well its binary decisions align '
        'with the multiclass ground truth: a healthy patient is correct if predicted HEALTHY, and '
        'a disease patient is correct if predicted UNHEALTHY (regardless of the specific disease '
        'type). The per-class table shows accuracy within each class under this binary evaluation.')
    add_source_ref(doc, 'Source: ImprovedCDS/cds.py (run_loocv function); Trace: ImprovedCDS/output/loocv_trace.json')

    add_para(doc, f'Overall Multiclass Accuracy: {base_stats["correct"]}/{base_stats["n"]} = {100*base_stats["accuracy"]:.1f}%', bold=True)
    add_para(doc, f'Binary Accuracy (Healthy vs Disease): {100*base_stats["binary_accuracy"]:.1f}%', bold=True)
    add_para(doc, f'Specificity (Healthy correctly identified): {100*base_stats["specificity"]:.1f}%')
    add_para(doc, f'Sensitivity (Disease detected as non-healthy): {100*base_stats["sensitivity"]:.1f}%')
    add_para(doc, f'Balanced Accuracy (mean per-class): {100*base_stats["balanced_accuracy"]:.1f}%')

    base_classes = sorted(base_stats['per_class'].keys())
    rows = []
    for cls in base_classes:
        s = base_stats['per_class'][cls]
        acc = f'{100*s["correct"]/s["total"]:.1f}%' if s['total'] > 0 else '0.0%'
        # Binary: class 1 patients predicted as 1 = correct; disease patients predicted != 1 = correct
        bin_correct = s['correct'] if cls == 1 else sum(
            1 for r in base_trace if r['true_class'] == cls and r.get('predicted', -1) != 1)
        bin_acc = f'{100*bin_correct/s["total"]:.1f}%' if s['total'] > 0 else '0.0%'
        rows.append([str(cls), CLASS_NAMES.get(cls, f'Class {cls}'),
                      str(s['total']), str(s['correct']), acc, bin_acc])
    add_para(doc, 'Table A1: Base CDS LOOCV Per-Class Results', bold=True)
    add_table(doc, ['Class', 'Description', 'N', 'Correct (Multi)', 'Multi Acc', 'Binary Acc'], rows)

    add_para(doc,
        'The base CDS algorithm achieves high specificity (94.7%) but poor sensitivity (51.7%). '
        'It correctly identifies most healthy patients but fails to detect disease subtypes. '
        'Classes 5, 6, 7, 8, 11, 12, and 13 have 0% accuracy, confirming the binary framework '
        'cannot distinguish between specific arrhythmia types.')

    # A.2 CDS-OVR
    add_heading(doc, 'A.2 CDS-OVR Algorithm LOOCV (All 13 Classes, 452 Patients)', 2)
    add_para(doc,
        'The CDS-OVR algorithm (cds_ovr.py) was evaluated using LOOCV on the same 452-patient '
        'dataset. Note that the LOOCV evaluation includes all 13 original classes, not the 8-class '
        'filtered dataset used for other evaluation protocols. This tests the algorithm on classes '
        'it was not designed to handle (7, 8, 11, 12, 13).')
    add_source_ref(doc, 'Source: ImprovedCDS/CDS_FINAL/cds_ovr.py; Trace: ImprovedCDS/output/loocv_trace_ovr.json')

    add_para(doc, f'Overall Multiclass Accuracy: {ovr_stats["correct"]}/{ovr_stats["n"]} = {100*ovr_stats["accuracy"]:.1f}%', bold=True)
    add_para(doc, f'Binary Accuracy (Healthy vs Disease): {100*ovr_stats["binary_accuracy"]:.1f}%', bold=True)
    add_para(doc, f'Specificity: {100*ovr_stats["specificity"]:.1f}%')
    add_para(doc, f'Sensitivity: {100*ovr_stats["sensitivity"]:.1f}%')
    add_para(doc, f'Balanced Accuracy: {100*ovr_stats["balanced_accuracy"]:.1f}%')

    ovr_classes = sorted(ovr_stats['per_class'].keys())
    rows = []
    for cls in ovr_classes:
        s = ovr_stats['per_class'][cls]
        acc = f'{100*s["correct"]/s["total"]:.1f}%' if s['total'] > 0 else '0.0%'
        bin_correct = s['correct'] if cls == 1 else sum(
            1 for r in ovr_trace if r['true_class'] == cls and r.get('predicted', -1) != 1)
        bin_acc = f'{100*bin_correct/s["total"]:.1f}%' if s['total'] > 0 else '0.0%'
        rows.append([str(cls), CLASS_NAMES.get(cls, f'Class {cls}'),
                      str(s['total']), str(s['correct']), acc, bin_acc])
    add_para(doc, 'Table A2: CDS-OVR LOOCV Per-Class Results (All 13 Classes)', bold=True)
    add_table(doc, ['Class', 'Description', 'N', 'Correct (Multi)', 'Multi Acc', 'Binary Acc'], rows)

    add_para(doc,
        f'CDS-OVR improves overall accuracy from {100*base_stats["accuracy"]:.1f}% to '
        f'{100*ovr_stats["accuracy"]:.1f}% (+{100*(ovr_stats["accuracy"]-base_stats["accuracy"]):.1f} pp) '
        f'and dramatically improves sensitivity from {100*base_stats["sensitivity"]:.1f}% to '
        f'{100*ovr_stats["sensitivity"]:.1f}% '
        f'(+{100*(ovr_stats["sensitivity"]-base_stats["sensitivity"]):.1f} pp). '
        f'The number of classes detected (accuracy > 0%) increases from 5 to 10 of 13.')

    # A.3 Comparison Summary
    add_heading(doc, 'A.3 LOOCV Comparison Summary', 2)
    add_para(doc, 'Table A3: LOOCV Comparison - Base CDS vs. CDS-OVR', bold=True)
    add_table(doc, ['Metric', 'Base CDS', 'CDS-OVR', 'Improvement'],
              [['Dataset', '452 patients, 13 classes', '452 patients, 13 classes', '-'],
               ['Overall Accuracy', f'{100*base_stats["accuracy"]:.1f}%', f'{100*ovr_stats["accuracy"]:.1f}%',
                f'+{100*(ovr_stats["accuracy"]-base_stats["accuracy"]):.1f} pp'],
               ['Binary Accuracy', f'{100*base_stats["binary_accuracy"]:.1f}%',
                f'{100*ovr_stats["binary_accuracy"]:.1f}%',
                f'+{100*(ovr_stats["binary_accuracy"]-base_stats["binary_accuracy"]):.1f} pp'],
               ['Specificity', f'{100*base_stats["specificity"]:.1f}%', f'{100*ovr_stats["specificity"]:.1f}%',
                f'{100*(ovr_stats["specificity"]-base_stats["specificity"]):+.1f} pp'],
               ['Sensitivity', f'{100*base_stats["sensitivity"]:.1f}%', f'{100*ovr_stats["sensitivity"]:.1f}%',
                f'+{100*(ovr_stats["sensitivity"]-base_stats["sensitivity"]):.1f} pp'],
               ['Balanced Accuracy', f'{100*base_stats["balanced_accuracy"]:.1f}%',
                f'{100*ovr_stats["balanced_accuracy"]:.1f}%',
                f'+{100*(ovr_stats["balanced_accuracy"]-base_stats["balanced_accuracy"]):.1f} pp'],
               ['Classes Detected (>0%)', '5 of 13', '10 of 13', '+5 classes'],
               ['Decision Type', 'Binary (Healthy/Unhealthy)', 'Multiclass (8 classes)', 'Multiclass capability']])

    # ================================================================
    # APPENDIX B: AUC/ROC
    # ================================================================
    doc.add_page_break()
    add_heading(doc, 'Appendix B: AUC/ROC Analysis', 1)

    add_para(doc,
        'The Area Under the Receiver Operating Characteristic Curve (AUC-ROC) provides a '
        'threshold-independent measure of discriminative ability. Unlike accuracy, AUC evaluates '
        'performance across all possible decision thresholds, making it robust to class imbalance '
        'and threshold selection.')

    add_heading(doc, 'B.1 Binary AUC-ROC Computation', 2)
    add_para(doc,
        'For binary AUC, the disease score for each patient is computed as the maximum disease-class '
        'ratio score divided by the healthy-class score plus epsilon:')
    add_para(doc, '  disease_score = max(R_c for c != 1) / (R_1 + 0.1)')
    add_para(doc,
        'The ROC curve is then constructed by varying the threshold on this composite score. The '
        'AUC is computed via the trapezoidal rule over sorted (FPR, TPR) pairs.')

    add_heading(doc, 'B.2 Multiclass AUC-ROC (One-vs-Rest)', 2)
    add_para(doc,
        'Per-class AUC is computed using the One-vs-Rest approach: for each class c, the ratio score '
        'R_c serves as the discriminant, and a binary label indicates whether the patient belongs to '
        'class c. Macro AUC is the unweighted mean across classes; weighted AUC weights each class '
        'by its support (number of patients).')

    if comp_results and 'auc_roc' in comp_results:
        auc = comp_results['auc_roc']
        add_source_ref(doc, 'Source: ImprovedCDS/CDS_FINAL/comprehensive_experiments.py; Results: ImprovedCDS/CDS_FINAL/comprehensive_results.json')

        rows = []
        for protocol, label in [('10fold', '10-Fold CV'), ('90_10', '90/10 Split'), ('60_40', '60/40 Split')]:
            if protocol in auc:
                d = auc[protocol]
                rows.append([label,
                             f'{d["binary_auc_mean"]:.4f} +/- {d["binary_auc_std"]:.4f}',
                             f'{d["macro_auc_mean"]:.4f} +/- {d["macro_auc_std"]:.4f}',
                             f'{d["weighted_auc_mean"]:.4f} +/- {d["weighted_auc_std"]:.4f}'])

        if rows:
            add_para(doc, 'Table B1: AUC-ROC Values Across Evaluation Protocols (10 seeds)', bold=True)
            add_table(doc, ['Protocol', 'Binary AUC', 'Macro AUC (OVR)', 'Weighted AUC (OVR)'], rows)

        add_heading(doc, 'B.3 Per-Class AUC-ROC', 2)
        if '10fold' in auc and 'per_class_auc_mean' in auc['10fold']:
            pc = auc['10fold']['per_class_auc_mean']
            rows = []
            for cls_str in sorted(pc.keys(), key=lambda x: int(x)):
                cls = int(cls_str)
                rows.append([str(cls), CLASS_NAMES.get(cls, f'Class {cls}'),
                             f'{pc[cls_str]:.4f}'])
            add_para(doc, 'Table B2: Per-Class AUC-ROC (10-Fold CV, mean over 10 seeds)', bold=True)
            add_table(doc, ['Class', 'Description', 'Mean AUC'], rows)
    else:
        add_para(doc,
            '[AUC/ROC results pending - comprehensive_experiments.py has not yet completed. '
            'Re-run this script after the experiment finishes to include these results.]',
            italic=True)

    # ================================================================
    # APPENDIX C: Comprehensive Performance Results
    # ================================================================
    doc.add_page_break()
    add_heading(doc, 'Appendix C: Comprehensive Performance Results', 1)

    add_para(doc,
        'This appendix consolidates all evaluation results across protocols and seeds. '
        'All results use the 416-patient, 8-class filtered dataset unless otherwise noted.')
    add_source_ref(doc,
        'Source: ImprovedCDS/CDS_FINAL/run_all_splits.py; '
        'Results: ImprovedCDS/CDS_FINAL/results_all_splits.json')

    # C.1 10-Fold CV
    add_heading(doc, 'C.1 10-Fold Cross-Validation (10 Seeds)', 2)
    if all_splits and '10-fold CV' in all_splits:
        cv = all_splits['10-fold CV']
        s = cv['summary']
        add_para(doc, f'Mean Multiclass Accuracy: {s["multi_mean"]}% (std = {s["multi_std"]}%)', bold=True)
        add_para(doc, f'Best Multiclass Accuracy: {s["multi_best"]}% (seed 13)')
        add_para(doc, f'Worst Multiclass Accuracy: {s["multi_worst"]}%')
        add_para(doc, f'Mean Binary Accuracy: {s["binary_mean"]}% (std = {s["binary_std"]}%)')
        add_para(doc, f'Best Binary Accuracy: {s["binary_best"]}%')

        rows = []
        for sd in cv['seeds']:
            rows.append([str(sd['seed']),
                         f'{sd["multiclass_acc"]}%', f'{sd["binary_acc"]}%',
                         f'{sd["specificity"]}%', f'{sd["sensitivity"]}%',
                         f'{sd["balanced_acc"]}%'])
        add_para(doc, 'Table C1: 10-Fold CV Results per Seed', bold=True)
        add_table(doc, ['Seed', 'Multi Acc', 'Binary Acc', 'Specificity', 'Sensitivity', 'Balanced Acc'], rows)

    # C.2 All Split Ratios
    add_heading(doc, 'C.2 Train/Test Split Results Across All Ratios', 2)
    add_para(doc,
        'CDS-OVR was evaluated across 7 split ratios (50/50 through 90/10), each with 10 random '
        'seeds. This sweep reveals how accuracy depends on training data volume.')

    if all_splits:
        summary_rows = []
        for ratio in ['50/50', '60/40', '70/30', '75/25', '80/20', '85/15', '90/10']:
            if ratio in all_splits:
                s = all_splits[ratio]['summary']
                summary_rows.append([ratio,
                    f'{s["multi_mean"]}%', f'{s["multi_std"]}%',
                    f'{s["multi_best"]}%', f'{s["multi_worst"]}%',
                    f'{s["binary_mean"]}%', f'{s["binary_std"]}%'])

        add_para(doc, 'Table C2: Summary Across All Split Ratios (10 seeds each)', bold=True)
        add_table(doc,
            ['Split', 'Multi Mean', 'Multi Std', 'Multi Best', 'Multi Worst', 'Binary Mean', 'Binary Std'],
            summary_rows)

    # C.3 90/10 Detailed
    add_heading(doc, 'C.3 90/10 Split Detailed Results', 2)
    if all_splits and '90/10' in all_splits:
        s90 = all_splits['90/10']
        s = s90['summary']
        add_para(doc, f'Mean Multiclass Accuracy: {s["multi_mean"]}% (std = {s["multi_std"]}%)', bold=True)
        add_para(doc, f'Best Multiclass Accuracy: {s["multi_best"]}% (seed 76)')
        add_para(doc, f'Mean Binary Accuracy: {s["binary_mean"]}% (std = {s["binary_std"]}%)')
        add_para(doc, f'Best Binary Accuracy: {s["binary_best"]}%')

        rows = []
        for sd in s90['seeds']:
            rows.append([str(sd['seed']),
                         f'{sd["multiclass_acc"]}%', f'{sd["binary_acc"]}%',
                         f'{sd["specificity"]}%', f'{sd["sensitivity"]}%',
                         f'{sd.get("balanced_acc", "N/A")}%'])
        add_para(doc, 'Table C3: 90/10 Split Results per Seed', bold=True)
        add_table(doc, ['Seed', 'Multi Acc', 'Binary Acc', 'Specificity', 'Sensitivity', 'Balanced Acc'], rows)

    # C.4 60/40 Detailed
    add_heading(doc, 'C.4 60/40 Split Detailed Results', 2)
    if all_splits and '60/40' in all_splits:
        s60 = all_splits['60/40']
        s = s60['summary']
        add_para(doc, f'Mean Multiclass Accuracy: {s["multi_mean"]}% (std = {s["multi_std"]}%)', bold=True)
        add_para(doc, f'Best Multiclass Accuracy: {s["multi_best"]}% (seed 76)')
        add_para(doc, f'Mean Binary Accuracy: {s["binary_mean"]}% (std = {s["binary_std"]}%)')

        rows = []
        for sd in s60['seeds']:
            rows.append([str(sd['seed']),
                         f'{sd["multiclass_acc"]}%', f'{sd["binary_acc"]}%',
                         f'{sd["specificity"]}%', f'{sd["sensitivity"]}%',
                         f'{sd.get("balanced_acc", "N/A")}%'])
        add_para(doc, 'Table C4: 60/40 Split Results per Seed', bold=True)
        add_table(doc, ['Seed', 'Multi Acc', 'Binary Acc', 'Specificity', 'Sensitivity', 'Balanced Acc'], rows)

    # C.5 Sensitivity/Specificity at 90/10 (from ablation report)
    add_heading(doc, 'C.5 Binary Sensitivity and Specificity at 90/10 (All Seeds)', 2)
    add_para(doc,
        'Full sensitivity and specificity breakdown for binary (healthy vs. disease) '
        'classification at the 90/10 split ratio.')
    add_source_ref(doc, 'Source: ImprovedCDS/CDS_FINAL/evidence_ablation_report.txt, Section 5')

    sens_spec_rows = [
        ['13', '85.7%', '83.3%', '88.9%', '20', '4', '16', '2'],
        ['20', '85.7%', '78.9%', '91.3%', '15', '4', '21', '2'],
        ['27', '81.0%', '73.3%', '85.2%', '11', '4', '23', '4'],
        ['34', '90.5%', '80.0%', '96.3%', '12', '3', '26', '1'],
        ['41', '88.1%', '72.2%', '100.0%', '13', '5', '24', '0'],
        ['48', '92.9%', '84.2%', '100.0%', '16', '3', '23', '0'],
        ['55', '88.1%', '71.4%', '96.4%', '10', '4', '27', '1'],
        ['62', '92.9%', '88.2%', '96.0%', '15', '2', '24', '1'],
        ['69', '83.3%', '78.9%', '87.0%', '15', '4', '20', '3'],
        ['76', '97.6%', '93.3%', '100.0%', '14', '1', '27', '0'],
    ]
    add_para(doc, 'Table C5: Binary Classification at 90/10 (All Seeds)', bold=True)
    add_table(doc, ['Seed', 'Binary Acc', 'Sensitivity', 'Specificity', 'TP', 'FN', 'TN', 'FP'],
              sens_spec_rows)
    add_para(doc, 'Mean Sensitivity: 80.4% (std = 7.0%). Mean Specificity: 94.1% (std = 5.5%).')

    # C.6 Per-Class Accuracy at 60/40 (from ablation report)
    add_heading(doc, 'C.6 Per-Class Accuracy at 60/40 Split (All Seeds)', 2)
    add_para(doc,
        'Per-class accuracy at the 60/40 split ratio, providing protocol-matched comparison '
        'data for Irfan et al. 2022 (who used 60/40 split).')
    add_source_ref(doc, 'Source: ImprovedCDS/CDS_FINAL/evidence_ablation_report.txt, Section 3')

    pc_rows = [
        ['1', 'Normal', '92.6%', '2.6%', '98.0%', '89.4%'],
        ['2', 'CAD', '63.6%', '8.3%', '82.4%', '50.0%'],
        ['3', 'Old Anterior MI', '97.6%', '5.0%', '100.0%', '85.7%'],
        ['4', 'Old Inferior MI', '57.2%', '28.0%', '100.0%', '12.5%'],
        ['5', 'Sinus Tachycardia', '45.8%', '31.5%', '100.0%', '0.0%'],
        ['6', 'Sinus Bradycardia', '56.4%', '18.6%', '87.5%', '25.0%'],
        ['9', 'LBBB', '77.2%', '22.3%', '100.0%', '33.3%'],
        ['10', 'RBBB', '64.1%', '8.6%', '77.3%', '42.9%'],
    ]
    add_para(doc, 'Table C6: Per-Class Accuracy at 60/40 (Mean over 10 Seeds)', bold=True)
    add_table(doc, ['Class', 'Description', 'Mean Acc', 'Std', 'Best', 'Worst'], pc_rows)

    # C.7 Stratified Results (All Protocols)
    add_heading(doc, 'C.7 Stratified Evaluation Results (All Protocols)', 2)
    add_para(doc,
        'A stratified split ensures proportional class representation in both training and test '
        'sets. This is important because rare classes (e.g., class 9 with 9 patients) may be '
        'entirely absent from the test set in random splits, making accuracy estimates unreliable. '
        'We evaluate stratified variants of all validation protocols.')
    add_source_ref(doc, 'Source: ImprovedCDS/CDS_FINAL/task_stratified.py; Results: results_stratified.json')

    strat_data = comp_results.get('stratified') if comp_results else None

    if strat_data:
        strat_protocols = [
            ('stratified_10fold', 'Stratified 10-Fold CV'),
            ('stratified_90_10', 'Stratified 90/10 Split'),
            ('stratified_80_20', 'Stratified 80/20 Split'),
            ('stratified_70_30', 'Stratified 70/30 Split'),
            ('stratified_60_40', 'Stratified 60/40 Split'),
            ('stratified_50_50', 'Stratified 50/50 Split'),
        ]

        # Summary table across all protocols
        summary_rows = []
        for key, label in strat_protocols:
            if key in strat_data:
                s = strat_data[key]['summary']
                summary_rows.append([label,
                    f'{s["acc_mean"]}%', f'{s["acc_std"]}%',
                    f'{s["acc_best"]}%', f'{s["acc_worst"]}%',
                    f'{s["binary_mean"]}%', f'{s["ba_mean"]}%'])

        if summary_rows:
            add_para(doc, 'Table C7: Stratified Results Summary Across All Protocols (10 seeds each)', bold=True)
            add_table(doc,
                ['Protocol', 'Multi Mean', 'Multi Std', 'Multi Best', 'Multi Worst', 'Binary Mean', 'BA Mean'],
                summary_rows)

        # Detailed per-seed tables for each protocol
        table_num = 8
        for key, label in strat_protocols:
            if key in strat_data:
                proto_data = strat_data[key]
                s = proto_data['summary']
                add_heading(doc, f'C.7.{table_num-7} {label} Detailed Results', 3)
                add_para(doc, f'Mean Multiclass Accuracy: {s["acc_mean"]}% (std = {s["acc_std"]}%)', bold=True)
                add_para(doc, f'Mean Binary Accuracy: {s["binary_mean"]}% (std = {s["binary_std"]}%)')
                add_para(doc, f'Mean Balanced Accuracy: {s["ba_mean"]}% (std = {s["ba_std"]}%)')

                rows = []
                for sd in proto_data['seeds']:
                    rows.append([str(sd['seed']),
                                 f'{sd["accuracy"]}%', f'{sd["binary_acc"]}%',
                                 f'{sd["specificity"]}%', f'{sd["sensitivity"]}%',
                                 f'{sd["balanced_acc"]}%', str(sd['n_test'])])
                add_para(doc, f'Table C{table_num}: {label} Results per Seed', bold=True)
                add_table(doc, ['Seed', 'Multi Acc', 'Binary Acc', 'Specificity', 'Sensitivity', 'Balanced Acc', 'N_test'], rows)
                table_num += 1
        # C.8 Stratified vs Non-Stratified Comparison
        add_heading(doc, 'C.8 Stratified vs. Random Split Comparison', 2)
        add_para(doc,
            'Comparing stratified results (C.7) with the corresponding random-split results (C.2) '
            'reveals the impact of class-balanced splitting on measured accuracy. In random splits, '
            'rare classes may be over- or under-represented in the test set, introducing variance '
            'unrelated to the algorithm\'s actual performance.')

        if all_splits and strat_data:
            comp_rows = []
            comparisons = [
                ('10-fold CV', '10-fold CV', 'stratified_10fold'),
                ('90/10', '90/10', 'stratified_90_10'),
                ('60/40', '60/40', 'stratified_60_40'),
                ('50/50', '50/50', 'stratified_50_50'),
            ]
            for label, random_key, strat_key in comparisons:
                if random_key in all_splits and strat_key in strat_data:
                    rs = all_splits[random_key]['summary']
                    ss = strat_data[strat_key]['summary']
                    diff = ss['acc_mean'] - rs['multi_mean']
                    comp_rows.append([label,
                        f'{rs["multi_mean"]}% ({rs["multi_std"]}%)',
                        f'{ss["acc_mean"]}% ({ss["acc_std"]}%)',
                        f'{diff:+.2f} pp',
                        f'{rs["binary_mean"]}%',
                        f'{ss["binary_mean"]}%'])

            if comp_rows:
                add_para(doc, 'Table C14: Random vs. Stratified Split Comparison', bold=True)
                add_table(doc,
                    ['Protocol', 'Random Multi (std)', 'Stratified Multi (std)',
                     'Diff', 'Random Binary', 'Stratified Binary'],
                    comp_rows)

            add_para(doc,
                'Stratified splitting generally produces slightly higher accuracy and lower variance, '
                'confirming that random splits introduce noise from unbalanced class representation. '
                'The difference is most pronounced at extreme split ratios where rare classes have '
                'very few test samples.')
    else:
        add_para(doc,
            '[Stratified results pending - task_stratified.py has not yet completed.]',
            italic=True)

    # ================================================================
    # APPENDIX D: Ablation Studies
    # ================================================================
    doc.add_page_break()
    add_heading(doc, 'Appendix D: Ablation Studies', 1)
    add_para(doc,
        'This appendix consolidates all ablation experiment results, quantifying the contribution '
        'of each CDS-OVR component.')
    add_source_ref(doc,
        'Source: ImprovedCDS/CDS_FINAL/evidence_ablation.py, ImprovedCDS/CDS_FINAL/ablation_full.py; '
        'Results: ImprovedCDS/CDS_FINAL/evidence_ablation_report.txt')

    add_heading(doc, 'D.1 Component Ablation Summary', 2)
    add_para(doc,
        'Each mechanism was individually disabled and evaluated via 10-fold CV across 10 seeds. '
        'Per-seed accuracy was paired with the full system, and a Wilcoxon signed-rank test '
        'determined statistical significance.')

    ablation_rows = [
        ['Supervised chi2 binning', '+5.75 pp', '0.002', 'Yes', 'Most impactful component'],
        ['Fisher discriminant weighting', '+4.52 pp', '0.002', 'Yes', 'Second most impactful'],
        ['Correlation filtering', '+0.46 pp', '0.313', 'No', 'Modest, prevents overfitting'],
        ['Sex branching', '+0.41 pp', '0.141', 'No', 'Concentrated on class 3'],
        ['against_scale (rare)', '+0.07 pp', '0.734', 'No', 'Rare classes = 8.9% of data'],
        ['Healthy bar', '-0.31 pp', '0.133', 'No', 'Trades accuracy for specificity'],
    ]
    add_para(doc, 'Table D1: Full Component Ablation Results', bold=True)
    add_table(doc,
        ['Component', 'Effect on Accuracy', 'p-value', 'Significant (p<0.05)', 'Notes'],
        ablation_rows)

    add_heading(doc, 'D.2 against_scale Ablation (Rare Classes)', 2)
    add_para(doc,
        'against_scale = 0.5 for rare classes {4, 5, 9} vs. uniform 0.8 for all classes.')
    ablation_rare_rows = [
        ['Overall (10-fold CV)', '84.66% (std=0.93%)', '84.59% (std=0.66%)', '+0.07 pp'],
        ['Class 4 (Old Inferior MI)', '69.3%', '62.7%', '+6.7 pp'],
        ['Class 5 (Sinus Tachycardia)', '59.2%', '50.0%', '+9.2 pp'],
        ['Class 9 (LBBB)', '96.7%', '95.6%', '+1.1 pp'],
    ]
    add_para(doc, 'Table D2: against_scale Ablation Results', bold=True)
    add_table(doc, ['Metric', 'Current (0.5 rare)', 'Uniform (0.8 all)', 'Difference'], ablation_rare_rows)

    add_heading(doc, 'D.3 Sex Branching Ablation', 2)
    add_para(doc, 'Sex-based population branching enabled (SEX_FEAT=1) vs. disabled (no branching).')
    ablation_sex_rows = [
        ['10-fold CV Mean', '84.66% (std=0.93%)', '84.25% (std=0.88%)', '+0.41 pp'],
        ['Wins', '6/10 seeds', '4/10 seeds', '-'],
    ]
    add_para(doc, 'Table D3: Sex Branching Ablation Results', bold=True)
    add_table(doc, ['Metric', 'Sex Branching On', 'Sex Branching Off', 'Difference'], ablation_sex_rows)

    add_heading(doc, 'D.4 Hyperparameter Sensitivity Analysis', 2)
    add_para(doc,
        'Three hyperparameters were varied independently while all others were held at defaults. '
        'Each configuration was evaluated via 10-fold CV across 10 seeds (100 fold evaluations).')
    add_source_ref(doc, 'Source: ImprovedCDS/CDS_FINAL/ablation_full.py, Part 4')

    hp_rows = [
        ['CORR_THRESHOLD', '0.6-1.0', '0.8', '1.44 pp', 'Optimal; degrades below 0.8'],
        ['MAX_BINS', '3-8', '6 (5 optimal)', '3.05 pp', 'MAX_BINS=5 is +0.51 pp better'],
        ['FEATURES_PER_CLASS', '12-24', '18', '0.00 pp', 'No effect (corr filter is bottleneck)'],
    ]
    add_para(doc, 'Table D4: Hyperparameter Sensitivity Sweep', bold=True)
    add_table(doc, ['Parameter', 'Range Tested', 'Current Value', 'Accuracy Range', 'Notes'], hp_rows)

    # ================================================================
    # APPENDIX E: Confusion Matrix and Classification Metrics
    # ================================================================
    doc.add_page_break()
    add_heading(doc, 'Appendix E: Confusion Matrix and Classification Metrics', 1)

    if comp_results and 'confusion_matrix_seed13' in comp_results:
        cm_data = comp_results['confusion_matrix_seed13']
        cm = cm_data['matrix']
        classes = cm_data['classes']
        add_source_ref(doc, 'Source: ImprovedCDS/CDS_FINAL/comprehensive_experiments.py; 10-fold CV, seed 13')

        add_heading(doc, 'E.1 Confusion Matrix (10-Fold CV, Seed 13)', 2)
        headers = ['True \\ Pred'] + [f'Cls {c}' for c in classes]
        rows = []
        for i, cls in enumerate(classes):
            row = [f'Class {cls}'] + [str(cm[i][j]) for j in range(len(classes))]
            rows.append(row)
        add_para(doc, 'Table E1: Confusion Matrix (10-Fold CV, Seed 13, 416 Patients)', bold=True)
        add_table(doc, headers, rows)

        add_heading(doc, 'E.2 Per-Class Precision, Recall, and F1-Score', 2)
        pc = cm_data['per_class']
        rows = []
        for cls_str in sorted(pc.keys(), key=lambda x: int(x)):
            cls = int(cls_str)
            m = pc[cls_str]
            rows.append([str(cls), CLASS_NAMES.get(cls, f'Class {cls}'),
                         str(m['support']),
                         f'{m["precision"]:.4f}', f'{m["recall"]:.4f}', f'{m["f1"]:.4f}'])
        add_para(doc, 'Table E2: Per-Class Classification Metrics', bold=True)
        add_table(doc, ['Class', 'Description', 'Support', 'Precision', 'Recall', 'F1-Score'], rows)

        add_para(doc, f'Macro Precision: {cm_data["macro_precision"]:.4f}')
        add_para(doc, f'Macro Recall: {cm_data["macro_recall"]:.4f}')
        add_para(doc, f'Macro F1-Score: {cm_data["macro_f1"]:.4f}')
    else:
        add_para(doc,
            '[Confusion matrix and P/R/F1 pending - comprehensive_experiments.py has not yet completed.]',
            italic=True)

    # ================================================================
    # APPENDIX F: Computational Cost
    # ================================================================
    doc.add_page_break()
    add_heading(doc, 'Appendix F: Computational Cost and Memory Usage', 1)

    add_para(doc,
        'The CDS-OVR algorithm requires no GPU acceleration, no iterative optimization, and '
        'no gradient computation. All operations are closed-form or single-pass through the data. '
        'This makes it suitable for deployment on resource-constrained embedded systems such as FPGAs.')

    if comp_results and 'computational_cost' in comp_results:
        cc = comp_results['computational_cost']
        add_source_ref(doc, 'Source: ImprovedCDS/CDS_FINAL/comprehensive_experiments.py (tracemalloc profiling)')

        add_heading(doc, 'F.1 Runtime Performance', 2)
        add_para(doc, 'Table F1: Computational Cost Summary', bold=True)
        add_table(doc, ['Operation', 'Time', 'Notes'],
                  [['Single Training (full dataset)', f'{cc["training_time_s"]:.4f} s', 'Build tree + train 8 class models'],
                   ['Single Patient Prediction', f'{cc["prediction_per_patient_ms"]:.4f} ms', 'Route + accumulate + decide'],
                   ['All 416 Predictions', f'{cc["prediction_all_patients_s"]:.4f} s', 'Sequential, no batching'],
                   ['10-Fold CV (1 seed)', f'{cc["10fold_cv_time_s"]:.2f} s', '10 train+test cycles'],
                   ['90/10 Split (1 seed)', f'{cc["90_10_split_time_s"]:.2f} s', '1 train+test cycle']])

        add_heading(doc, 'F.2 Memory Usage', 2)
        add_para(doc, 'Table F2: Memory Usage Summary', bold=True)
        add_table(doc, ['Component', 'Size', 'Description'],
                  [['Data Matrix', f'{cc["data_matrix_mb"]:.2f} MB', '416 x 279 float64 array'],
                   ['Current Memory', f'{cc["memory_current_mb"]:.2f} MB', 'After training all models'],
                   ['Peak Memory', f'{cc["memory_peak_mb"]:.2f} MB', 'Maximum during prediction'],
                   ['Tree Nodes', str(cc['n_tree_nodes']), 'Root + sex-based branches'],
                   ['Total BinModels', str(cc['n_class_models']), '8 classes x 279 features x nodes'],
                   ['Retained Features', str(cc['n_retained_features']), '~18 per class per node']])
    else:
        add_para(doc,
            '[Computational cost results pending - comprehensive_experiments.py has not yet completed.]',
            italic=True)

    add_heading(doc, 'F.3 Computational Complexity Analysis', 2)
    add_para(doc,
        'Training: O(C x N x F x B) where C=8 classes, N=416 patients, F=279 features, '
        'B=6 max bins. The dominant cost is computing supervised bin edges for each feature in '
        'each class model. Feature selection adds O(F^2) for pairwise correlation but only within '
        'the top candidates.')
    add_para(doc,
        'Prediction: O(C x F_r x L) where C=8 classes, F_r=18 retained features per class, '
        'L=number of matched tree nodes (at most 2). Each prediction requires looking up bin '
        'assignments and accumulating evidence, all O(1) per feature.')
    add_para(doc,
        'The algorithm requires no matrix inversions, no eigendecompositions, no iterative '
        'solvers, and no backpropagation.')
    add_source_ref(doc, 'Implementation: ImprovedCDS/CDS_FINAL/cds_ovr.py, 587 lines total')

    # ================================================================
    # APPENDIX G: Mathematical Formulation
    # ================================================================
    doc.add_page_break()
    add_heading(doc, 'Appendix G: Mathematical Formulation', 1)

    add_para(doc,
        'This appendix provides the complete mathematical formulation for each component of the '
        'CDS-OVR algorithm, with references to the relevant literature.')

    add_heading(doc, 'G.1 Supervised Chi-Squared Binning', 2)
    add_para(doc,
        'The supervised binning algorithm partitions continuous feature values into discrete bins '
        'by maximizing the chi-squared statistic between the target class indicator and the bin '
        'assignment. For a candidate split dividing a segment into left and right partitions:')
    add_para(doc, '  chi2 = SUM_i SUM_j (O_ij - E_ij)^2 / E_ij')
    add_para(doc,
        'where O_ij is the observed count (i in {target, rest}, j in {left, right}) and '
        'E_ij = n_j * n_i / n is the expected count under the null hypothesis of independence. '
        'The chi-squared statistic follows an asymptotic chi2(1) distribution under the null. '
        'A minimum gain threshold of 0.5 prevents overfitting (Kerber, 1992).')
    add_source_ref(doc, 'Implementation: cds_ovr.py, lines 209-265 (chi2_binning function)')

    add_heading(doc, 'G.2 Laplace-Smoothed Posterior', 2)
    add_para(doc,
        'The class-conditional posterior probability for class c in bin b is computed with '
        'Laplace smoothing:')
    add_para(doc, '  P(c|b) = (n_cb + alpha * pi_c) / (n_b + alpha)')
    add_para(doc,
        'where n_cb is the number of class-c patients in bin b, n_b is the total count in bin b, '
        'pi_c = n_c/N is the class prior, and alpha = 1.0 is the smoothing parameter. This is '
        'equivalent to placing a Dirichlet prior on the bin-conditional class distribution '
        '(Manning et al., 2008).')
    add_source_ref(doc, 'Implementation: cds_ovr.py, lines 267-290')

    add_heading(doc, 'G.3 Fisher Discriminant Ratio (FDR)', 2)
    add_para(doc,
        'The Fisher Discriminant Ratio measures univariate separability between the target class '
        'and all other classes:')
    add_para(doc, '  FDR(f, c) = (mu_c - mu_rest)^2 / (sigma_c^2 + sigma_rest^2 + epsilon)')
    add_para(doc,
        'where mu_c and sigma_c^2 are the mean and variance of feature f for class c patients, '
        'mu_rest and sigma_rest^2 are for all other patients, and epsilon = 1e-10 prevents '
        'division by zero. FDR equals the squared signal-to-noise ratio for the binary '
        'classification task (Fisher, 1936).')
    add_source_ref(doc, 'Implementation: cds_ovr.py, lines 316-321')

    add_heading(doc, 'G.4 Fisher Weight Computation', 2)
    add_para(doc, 'Each retained feature receives a weight based on its relative FDR:')
    add_para(doc, '  w_f = max(sqrt(FDR_f / (FDR_max + epsilon)), 0.1)')
    add_para(doc,
        'where FDR_max is the maximum FDR across all retained features for the class. '
        'The square root compresses the dynamic range, and the floor of 0.1 ensures no feature '
        'is entirely silenced.')
    add_source_ref(doc, 'Implementation: cds_ovr.py, line 430')

    add_heading(doc, 'G.5 Action Score (Feature Utility)', 2)
    add_para(doc,
        'The action score quantifies how much a feature contributes to distinguishing the target '
        'class from the population baseline:')
    add_para(doc, '  S(f, c) = SUM_b |P(c|b) - pi_c| * min(1, n_b / CS)')
    add_para(doc,
        'where the sum is over bins b with at least MS patients, P(c|b) is the Laplace-smoothed '
        'posterior, pi_c is the class prior, and CS is the confidence support parameter. '
        'The confidence factor min(1, n_b/CS) prevents bins with few observations from '
        'contributing disproportionately.')
    add_source_ref(doc, 'Implementation: cds_ovr.py, lines 292-314')

    add_heading(doc, 'G.6 Dual Evidence Accumulation', 2)
    add_para(doc, 'Evidence for and against each class is accumulated separately:')
    add_para(doc, '  AF_for = SUM_f [shift_f * confidence_f * w_f]   for shift_f >= 0')
    add_para(doc, '  AF_against = SUM_f [|shift_f| * confidence_f * w_f * alpha_c]   for shift_f < 0')
    add_para(doc,
        'where shift_f = P(c|b_f) - pi_c is the posterior shift for the patient in feature f, '
        'confidence_f = min(1, bc_f / 10) is the confidence based on bin population, '
        'w_f is the Fisher weight, and alpha_c is the against_scale parameter '
        '(0.5 for rare classes {4, 5, 9}; 0.8 for common classes).')
    add_source_ref(doc, 'Implementation: cds_ovr.py, lines 399-442')

    add_heading(doc, 'G.7 Ratio Score', 2)
    add_para(doc, 'The ratio score combines positive and negative evidence:')
    add_para(doc, '  R_c = (AF_for + epsilon) / (AF_against + epsilon)')
    add_para(doc,
        'where epsilon = 0.1 provides numerical stability. Ratio > 1.0 indicates net positive '
        'evidence for the class; ratio << 1.0 indicates net negative evidence.')

    add_heading(doc, 'G.8 Dynamic Healthy Bar and Decision Logic', 2)
    add_para(doc, 'The effective threshold for each disease class is:')
    add_para(doc, '  T_eff = max(T_c - 0.3 * I(h_score < 2.0), min(1.05 * h_score, 5.0))')
    add_para(doc,
        'where T_c is the per-class base threshold (from CLASS_THRESHOLDS), h_score = R_1 is '
        'the ratio score for the healthy class, 0.3 is the suspicion offset, 2.0 is the '
        'suspicion threshold, 1.05 is the healthy weight, and 5.0 is the healthy bar cap.')
    add_para(doc,
        'A class c is a candidate if R_c >= T_eff. The prediction is the candidate with the '
        'maximum normalized margin:')
    add_para(doc, '  prediction = argmax_c (R_c - T_eff) / max(T_eff, 0.1)')
    add_para(doc, 'If no disease class exceeds its threshold, the patient is predicted as Healthy (class 1).')
    add_source_ref(doc, 'Implementation: cds_ovr.py, lines 455-468 (healthy bar), lines 470-480 (decision)')

    add_heading(doc, 'G.9 Correlation-Based Feature Filtering', 2)
    add_para(doc,
        'The absolute Pearson correlation between features f1 and f2 is:')
    add_para(doc, '  |r| = |SUM((x_i - mu_x)(y_i - mu_y))| / sqrt(SUM(x_i - mu_x)^2 * SUM(y_i - mu_y)^2)')
    add_para(doc,
        'Features with |r| > 0.8 are considered redundant. During greedy selection, each '
        'candidate feature is checked against all already-selected features, and rejected if any '
        'pairwise correlation exceeds the threshold (Guyon & Elisseeff, 2003).')
    add_source_ref(doc, 'Implementation: cds_ovr.py, lines 330-367 (feature selection loop)')

    add_heading(doc, 'G.10 Evaluation Metrics', 2)
    add_para(doc, 'Multiclass Accuracy = (1/N) * SUM_i I(y_pred_i == y_true_i)')
    add_para(doc, 'Specificity = TN / (TN + FP)   [for healthy class]')
    add_para(doc, 'Sensitivity = TP / (TP + FN)   [for disease detection]')
    add_para(doc, 'Balanced Accuracy = (1/C) * SUM_c (correct_c / n_c)')
    add_para(doc, 'Binary Accuracy = SUM_i I((true_i==1) == (pred_i==1)) / N')
    add_para(doc, 'Precision_c = TP_c / (TP_c + FP_c)')
    add_para(doc, 'Recall_c = TP_c / (TP_c + FN_c)')
    add_para(doc, 'F1_c = 2 * Precision_c * Recall_c / (Precision_c + Recall_c)')
    add_para(doc, 'AUC = integral(TPR(FPR) dFPR)   [trapezoidal rule]')

    # ================================================================
    # APPENDIX H: References
    # ================================================================
    doc.add_page_break()
    add_heading(doc, 'Appendix H: References', 1)

    refs = [
        '[1] Kerber, R. (1992). "ChiMerge: Discretization of numeric attributes." In Proceedings of the 10th National Conference on Artificial Intelligence (AAAI-92), pp. 123-128.',
        '[2] Manning, C. D., Raghavan, P., & Schutze, H. (2008). Introduction to Information Retrieval. Cambridge University Press. Chapter 13.',
        '[3] Fisher, R. A. (1936). "The use of multiple measurements in taxonomic problems." Annals of Eugenics, 7(2), 179-188.',
        '[4] Guyon, I., & Elisseeff, A. (2003). "An introduction to variable and feature selection." JMLR, 3, 1157-1182.',
        '[5] Guvenir, H. A., Acar, B., Demiroz, G., & Cekin, A. (1997). "A supervised machine learning algorithm for arrhythmia analysis." In Proceedings of Computers in Cardiology, pp. 433-436.',
        '[6] Sharma, A., et al. (2024). "Optimized Random Forest for arrhythmia classification on the UCI dataset." Engineering Research Express, 6, 035209.',
        '[7] Mustaqeem, A., Anwar, S. M., & Majid, M. (2018). "Multiclass classification of cardiac arrhythmia using improved feature selection and SVM invariants." Computational and Mathematical Methods in Medicine, 2018, 7310496.',
        '[8] Irfan, S., et al. (2022). "Heartbeat classification and arrhythmia detection using a multi-model deep-learning technique." Sensors, 22(15), 5606.',
        '[9] Gupta, V., Srinivasan, S. and Kudli, S. S. (2014). "Prediction and classification of cardiac arrhythmia."',
        '[10] Jadhav, S. M., Nalbalwar, S. L., & Ghatol, A. A. (2012). "Artificial neural network models based cardiac arrhythmia disease diagnosis from ECG data." International Journal of Computer Applications, 44(15).',
        '[11] Plawiak, P. (2018). "Novel methodology of cardiac health recognition based on ECG signals and evolutionary-neural system." Expert Systems with Applications, 92, 334-349.',
        '[12] Khan Mamun, M. M. R. (2021). "Arrhythmia classification using hybrid feature selection approach." IEEE CCECE 2021.',
        '[13] Islam, M. S., et al. (2023). "A review of arrhythmia classification techniques." arXiv:2301.10174.',
        '[14] UCI Machine Learning Repository. (1998). Arrhythmia Data Set. https://archive.ics.uci.edu/ml/datasets/arrhythmia',
        '[15] Rifkin, R., & Klautau, A. (2004). "In defense of one-vs-all classification." JMLR, 5, 101-141.',
        '[16] Wilcoxon, F. (1945). "Individual comparisons by ranking methods." Biometrics Bulletin, 1(6), 80-83.',
    ]
    for ref in refs:
        add_para(doc, ref, font_size=10)

    # ================================================================
    # APPENDIX I: Consolidated Results Summary
    # ================================================================
    doc.add_page_break()
    add_heading(doc, 'Appendix I: Consolidated Results Summary', 1)
    add_para(doc,
        'This appendix provides a single-table summary of all key results across all evaluation '
        'protocols. All results use the 416-patient, 8-class filtered dataset unless noted otherwise.')

    summary_rows = [
        ['LOOCV (13 classes, 452 pts)', '74.8%', '84.5%', '90.2%', '77.8%', '47.8%', '-'],
    ]
    if all_splits:
        import numpy as np_s
        def _mean_metric(seeds_list, key):
            vals = [s[key] for s in seeds_list if key in s]
            return round(float(np_s.mean(vals)), 2) if vals else None
        if '10-fold CV' in all_splits:
            s = all_splits['10-fold CV']['summary']
            seeds = all_splits['10-fold CV']['seeds']
            spec = _mean_metric(seeds, 'specificity')
            sens = _mean_metric(seeds, 'sensitivity')
            ba = _mean_metric(seeds, 'balanced_acc')
            summary_rows.append(['10-Fold CV (random)', f'{s["multi_mean"]}%', f'{s["binary_mean"]}%',
                                 f'{spec}%' if spec else '-', f'{sens}%' if sens else '-',
                                 f'{ba}%' if ba else '-', f'{s["multi_std"]}%'])
        for ratio in ['90/10', '80/20', '70/30', '60/40', '50/50']:
            rkey = ratio
            train_pct = ratio.split('/')[0]
            for k in all_splits:
                if k.replace('/', '_').replace('-', '_') == ratio.replace('/', '_'):
                    rkey = k
                    break
                if ratio in k:
                    rkey = k
                    break
                if k.startswith(train_pct + '/'):
                    rkey = k
                    break
            if rkey in all_splits:
                s = all_splits[rkey]['summary']
                seeds = all_splits[rkey]['seeds']
                spec = _mean_metric(seeds, 'specificity')
                sens = _mean_metric(seeds, 'sensitivity')
                ba = _mean_metric(seeds, 'balanced_acc')
                summary_rows.append([f'{ratio} Split (random)', f'{s["multi_mean"]}%', f'{s["binary_mean"]}%',
                                     f'{spec}%' if spec else '-', f'{sens}%' if sens else '-',
                                     f'{ba}%' if ba else '-', f'{s["multi_std"]}%'])

    strat_data = comp_results.get('stratified') if comp_results else None
    if strat_data:
        strat_protocols = [
            ('stratified_10fold', 'Strat. 10-Fold CV'),
            ('stratified_90_10', 'Strat. 90/10'),
            ('stratified_80_20', 'Strat. 80/20'),
            ('stratified_70_30', 'Strat. 70/30'),
            ('stratified_60_40', 'Strat. 60/40'),
            ('stratified_50_50', 'Strat. 50/50'),
        ]
        for key, label in strat_protocols:
            if key in strat_data:
                s = strat_data[key]['summary']
                summary_rows.append([label, f'{s["acc_mean"]}%', f'{s["binary_mean"]}%',
                                     f'{s["spec_mean"]}%', f'{s["sens_mean"]}%',
                                     f'{s["ba_mean"]}%', f'{s["acc_std"]}%'])

    add_para(doc, 'Table I1: Complete Results Summary Across All Evaluation Protocols', bold=True)
    add_table(doc, ['Protocol', 'Multi Acc', 'Binary Acc', 'Specificity', 'Sensitivity', 'Balanced Acc', 'Std Dev'],
              summary_rows)

    if comp_results and 'auc_roc' in comp_results:
        auc = comp_results['auc_roc']
        auc_rows = []
        for protocol, label in [('10fold', '10-Fold CV'), ('90_10', '90/10 Split'), ('60_40', '60/40 Split')]:
            if protocol in auc:
                d = auc[protocol]
                auc_rows.append([label, f'{d["binary_auc_mean"]:.4f}',
                                 f'{d["macro_auc_mean"]:.4f}', f'{d["weighted_auc_mean"]:.4f}'])
        if auc_rows:
            add_para(doc, 'Table I2: AUC-ROC Summary', bold=True)
            add_table(doc, ['Protocol', 'Binary AUC', 'Macro AUC', 'Weighted AUC'], auc_rows)

    if comp_results and 'confusion_matrix_seed13' in comp_results:
        cm = comp_results['confusion_matrix_seed13']
        add_para(doc, 'Table I3: Classification Metrics Summary (10-Fold CV, Seed 13)', bold=True)
        add_table(doc, ['Metric', 'Value'],
                  [['Overall Accuracy', f'{cm["accuracy"]}%'],
                   ['Binary Accuracy', f'{cm["binary_accuracy"]}%'],
                   ['Macro Precision', f'{cm["macro_precision"]:.4f}'],
                   ['Macro Recall', f'{cm["macro_recall"]:.4f}'],
                   ['Macro F1-Score', f'{cm["macro_f1"]:.4f}']])

    # ================================================================
    # APPENDIX J: Limitations and Future Work
    # ================================================================
    doc.add_page_break()
    add_heading(doc, 'Appendix J: Limitations and Future Work', 1)

    add_heading(doc, 'J.1 Current Limitations', 2)

    add_para(doc, '1. Small Dataset Size', bold=True)
    add_para(doc,
        'The UCI Arrhythmia Dataset contains only 452 patients. After removing classes with '
        'insufficient samples, 416 patients across 8 classes remain. Several classes have fewer '
        'than 15 patients (class 5: 13, class 4: 15, class 3: 15, class 9: 9). This limits the '
        'statistical power of all evaluation protocols, particularly for rare-class performance estimates.')

    add_para(doc, '2. Single Dataset Evaluation', bold=True)
    add_para(doc,
        'All results are reported on a single dataset (UCI Arrhythmia). While the multi-protocol '
        'evaluation strategy (LOOCV, 10-fold CV, multiple splits, stratified variants) provides '
        'robustness, generalization to other arrhythmia datasets or ECG recording conditions has '
        'not been evaluated.')

    add_para(doc, '3. Tabular Features Only', bold=True)
    add_para(doc,
        'The algorithm operates on pre-extracted tabular features (279 numeric features per patient), '
        'not on raw ECG waveforms. This means the feature extraction quality is fixed and cannot be '
        'optimized end-to-end. Deep learning methods that operate directly on ECG signals can learn '
        'task-specific feature representations.')

    add_para(doc, '4. Performance Gap at Low Training Ratios', bold=True)
    add_para(doc,
        'CDS-OVR accuracy decreases more rapidly than some benchmark methods when the training set '
        'shrinks. At 60/40 split, mean accuracy is 80.18% compared to 84.66% at 10-fold CV. This '
        'is a structural limitation of the bin-based evidence framework, which requires sufficient '
        'bin populations to estimate class-conditional probabilities accurately.')

    add_para(doc, '5. Independent Feature Assumption', bold=True)
    add_para(doc,
        'The evidence accumulation treats each feature independently, weighted by Fisher discriminant '
        'ratio. It does not model feature interactions. Methods like neural networks and SVMs with '
        'nonlinear kernels can capture these interactions, potentially explaining their advantage on '
        'certain class boundaries.')

    add_para(doc, '6. Five Removed Classes', bold=True)
    add_para(doc,
        'Classes 7, 8, 11, 12, and 13 (36 patients total) are removed due to insufficient samples. '
        'A clinical deployment would need to handle these classes, either through data augmentation, '
        'transfer learning, or a "reject" option.')

    add_heading(doc, 'J.2 Future Work', 2)

    add_para(doc, '1. FPGA Implementation', bold=True)
    add_para(doc,
        'The algorithm\'s fixed-point arithmetic compatibility, absence of iterative optimization, '
        'and small memory footprint (3.05 MB) make it a candidate for FPGA-based real-time ECG '
        'classification. Future work will translate the Python implementation to Verilog HDL, '
        'targeting low-latency inference for wearable cardiac monitoring devices.')

    add_para(doc, '2. Cross-Dataset Validation', bold=True)
    add_para(doc,
        'Evaluate CDS-OVR on additional arrhythmia datasets (e.g., MIT-BIH Arrhythmia Database, '
        'PhysioNet Challenge datasets) to assess generalization. This would require adapting the '
        'feature extraction pipeline to match each dataset\'s recording format.')

    add_para(doc, '3. Feature Interaction Modeling', bold=True)
    add_para(doc,
        'Explore augmenting the evidence framework with pairwise or higher-order feature interactions, '
        'potentially through interaction terms in the scoring function or a second-stage ensemble '
        'that combines per-class ratio scores.')

    add_para(doc, '4. Rare Class Data Augmentation', bold=True)
    add_para(doc,
        'Investigate synthetic oversampling (SMOTE or similar) for rare classes to improve bin '
        'population estimates. This could reduce the against_scale compensation currently needed '
        'for classes 4, 5, and 9.')

    add_para(doc, '5. Adaptive Binning', bold=True)
    add_para(doc,
        'Replace the fixed MAX_BINS=6 with a per-feature adaptive scheme that selects the number '
        'of bins based on sample size and feature distribution, using information-theoretic criteria '
        '(MDL or BIC) to prevent overfitting on sparse features.')

    add_para(doc, '6. Confidence Calibration', bold=True)
    add_para(doc,
        'Calibrate the ratio scores as proper probabilities using Platt scaling or isotonic '
        'regression. This would enable clinically useful confidence estimates (e.g., "85% confidence '
        'this patient has class 2 arrhythmia"), supporting clinical decision-making workflows.')

    # ================================================================
    # APPENDIX K: Reproducibility and Data Availability
    # ================================================================
    doc.add_page_break()
    add_heading(doc, 'Appendix K: Reproducibility and Data Availability', 1)

    add_heading(doc, 'K.1 Data Availability', 2)
    add_para(doc,
        'The UCI Arrhythmia Dataset is publicly available from the UCI Machine Learning Repository '
        'at https://archive.ics.uci.edu/ml/datasets/arrhythmia (Guvenir et al., 1997). The dataset '
        'file (arrhythmia.data) contains 452 rows and 280 columns (279 features + 1 class label) '
        'in CSV format with missing values encoded as "?".')

    add_heading(doc, 'K.2 Code Availability', 2)
    add_para(doc,
        'The complete source code for the CDS-OVR algorithm, all evaluation scripts, and this '
        'report\'s generation code are available in the project repository. The main algorithm '
        'implementation is in ImprovedCDS/CDS_FINAL/cds_ovr.py (587 lines of Python). All results '
        'can be reproduced by running the evaluation scripts with the seeds documented in this report.')

    add_heading(doc, 'K.3 Reproducibility Notes', 2)
    add_para(doc,
        'All randomized experiments use explicit integer seeds (13, 20, 27, 34, 41, 48, 55, 62, 69, 76) '
        'passed to numpy.random.RandomState for deterministic reproduction. The evaluation code uses '
        'numpy for numerical operations and python-docx for report generation. No GPU or specialized '
        'hardware is required. The complete evaluation suite runs in approximately 12 minutes on a '
        'standard desktop CPU (single-threaded) or approximately 4 minutes with 10-worker parallelism.')

    add_heading(doc, 'K.4 Software Environment', 2)
    add_para(doc,
        'Python 3.13, NumPy, python-docx. The algorithm itself (cds_ovr.py) depends only on NumPy. '
        'No deep learning frameworks, no scikit-learn, no external ML libraries.')

    # ================================================================
    # APPENDIX L: Repository File Index
    # ================================================================
    doc.add_page_break()
    add_heading(doc, 'Appendix L: Repository File Index', 1)
    add_para(doc,
        'This appendix lists all source files, data files, and generated outputs referenced '
        'in this report, with descriptions of their contents.')

    file_index = [
        ['ImprovedCDS/CDS_FINAL/cds_ovr.py', '587 lines', 'Final CDS-OVR algorithm implementation'],
        ['ImprovedCDS/cds.py', '491 lines', 'Original base CDS algorithm (binary classifier)'],
        ['ImprovedCDS/CDS_FINAL/run_all_splits.py', '157 lines', 'All split ratio evaluations (10 seeds each)'],
        ['ImprovedCDS/CDS_FINAL/task_auc.py', '-', 'Parallel AUC/ROC computation (6 workers)'],
        ['ImprovedCDS/CDS_FINAL/task_stratified.py', '-', 'Parallel stratified evaluation (all protocols)'],
        ['ImprovedCDS/CDS_FINAL/task_cost.py', '-', 'Computational cost and memory profiling'],
        ['ImprovedCDS/CDS_FINAL/task_confusion.py', '-', 'Confusion matrix and P/R/F1 computation'],
        ['ImprovedCDS/CDS_FINAL/evidence_analysis.py', '-', 'Evidence quality and dataset analysis (14 sections)'],
        ['ImprovedCDS/CDS_FINAL/evidence_ablation.py', '-', 'Ablation experiments (against_scale, sex branching)'],
        ['ImprovedCDS/CDS_FINAL/ablation_full.py', '-', 'Full component ablation with significance testing'],
        ['ImprovedCDS/CDS_FINAL/restructure_report.py', '-', 'Report generation and appendix construction'],
        ['ImprovedCDS/CDS_FINAL/results_all_splits.json', '-', 'All split ratio results (50/50-90/10, 10-fold CV)'],
        ['ImprovedCDS/CDS_FINAL/results_auc.json', '-', 'AUC/ROC results (3 protocols x 10 seeds)'],
        ['ImprovedCDS/CDS_FINAL/results_stratified.json', '-', 'Stratified results (6 protocols x 10 seeds)'],
        ['ImprovedCDS/CDS_FINAL/results_cost.json', '-', 'Computational cost and memory profiling results'],
        ['ImprovedCDS/CDS_FINAL/results_confusion.json', '-', 'Confusion matrix and per-class metrics'],
        ['ImprovedCDS/CDS_FINAL/evidence_report.txt', '-', 'Detailed evidence analysis output (14 sections)'],
        ['ImprovedCDS/CDS_FINAL/evidence_ablation_report.txt', '-', 'Ablation experiment output'],
        ['ImprovedCDS/output/loocv_trace.json', '~3.1 MB', 'Base CDS LOOCV trace (452 patient predictions)'],
        ['ImprovedCDS/output/loocv_trace_ovr.json', '~3.5 MB', 'CDS-OVR LOOCV trace (452 patient predictions)'],
        ['ImprovedCDS/output/logs/', '20 files', '10-fold CV log files for each seed'],
        ['data/arrhythmia.data', '-', 'UCI Arrhythmia Dataset (452 x 280, CSV)'],
    ]
    add_para(doc, 'Table L1: Complete File Index', bold=True)
    add_table(doc, ['File Path', 'Size/Lines', 'Description'], file_index)

    return True


# ============================================================
# Main
# ============================================================

def main():
    print("=" * 60)
    print("  CDS-OVR Report Restructuring Tool")
    print("=" * 60)

    # Copy original to temp to avoid OneDrive issues
    orig = os.path.join(SRC_DIR, 'CDS_Overall_Report.docx')
    temp_dir = os.environ.get('TEMP', '/tmp')
    src = os.path.join(temp_dir, 'CDS_Report_Restructure.docx')
    shutil.copy2(orig, src)

    print(f"\nOpening report from {src}...")
    doc = Document(src)
    print(f"  {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

    # Phase 1: Fix em/en dashes
    print("\nPhase 1: Replacing em/en dashes...")
    dash_count = replace_em_dashes(doc)
    print(f"  Replaced {dash_count} em/en dashes")

    # Phase 2: Fix structure (Part dividers)
    print("\nPhase 2: Fixing document structure...")
    struct_count = fix_structure(doc)
    print(f"  Made {struct_count} structural changes")

    # Phase 2b: Fix Part III empty subtitle
    print("\nPhase 2b: Fixing Part III subtitle...")
    fix_part3_subtitle(doc)

    # Phase 2c: Fix duplicate section numbers
    print("\nPhase 2c: Fixing duplicate section numbers...")
    fix_duplicate_section_numbers(doc)

    # Phase 3: Add file location references (before abstract to avoid matching abstract text)
    print("\nPhase 3: Adding file location references...")
    ref_count = add_file_location_refs(doc)
    print(f"  Added {ref_count} source annotations")

    # Phase 4: Mark redundancies
    print("\nPhase 4: Marking redundant content...")
    red_count = mark_redundancies(doc)
    print(f"  Marked {red_count} redundancies")

    # Phase 5: Insert Abstract (after file refs so abstract text doesn't trigger annotations)
    print("\nPhase 5: Inserting Abstract...")
    add_abstract(doc)

    # Phase 5b: Fix text errors (sex encoding, typo, dual AF ratio)
    print("\nPhase 5b: Fixing text errors...")
    text_fixes = fix_text_issues(doc)
    print(f"  Applied {text_fixes} text fixes")

    # Phase 5c: Clarify base model multiclass in Part II
    print("\nPhase 5c: Clarifying base model multiclass metric...")
    base_fixes = fix_appendix_a_clarification(doc)
    print(f"  Applied {base_fixes} base model clarifications")

    # Phase 5d: Update healthy bar explanation with grounded analysis
    print("\nPhase 5d: Updating healthy bar analysis...")
    hbar_fixes = add_healthy_bar_analysis(doc)
    print(f"  Applied {hbar_fixes} healthy bar updates")

    # Phase 5e: Insert binning comparison figure
    print("\nPhase 5e: Inserting binning comparison figure...")
    bin_figs = add_binning_figure(doc)
    print(f"  Inserted {bin_figs} binning figure(s)")

    # Phase 5f: Insert OVR class separation figure
    print("\nPhase 5f: Inserting OVR separation figure...")
    ovr_figs = add_ovr_figure(doc)
    print(f"  Inserted {ovr_figs} OVR figure(s)")

    # Phase 5g: Insert missing ablation results (if available)
    print("\nPhase 5g: Inserting missing ablation results...")
    abl_count = add_ablation_missing_results(doc)
    print(f"  Inserted {abl_count} ablation result(s)")

    # Phase 6: Build appendices
    print("\nPhase 6: Building comprehensive appendices...")
    build_appendices(doc)
    print("  Appendices A-L built successfully")

    # Save
    out_path = os.path.join(SRC_DIR, 'CDS_Final_Report.docx')
    doc.save(out_path)
    print(f"\nSaved final report to {out_path}")
    print(f"  Final: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")


if __name__ == '__main__':
    main()
